from django.shortcuts import render
from django.shortcuts import HttpResponse
from .forms import ParticipantForm, Breath_Collection_Form, Blood_Collection_Form, Lab_Processing_Form, Exposure_Form, Exposure_Form2, Exposure_Form3, Indirect_costs_form, Stai_y1_2_form, Veterans_rand_12_form, Mandatory_questionaire_form, SearchForm, CT_Scan_Form_T7, CT_Scan_Form_T6, CT_Scan_Form_T5, CT_Scan_Form_T4, CT_Scan_Form_T3, CT_Scan_Form_T2, CT_Scan_Form_T1, CT_Scan_Form_Baseline, Mandatory_questionaire_form_c, Mandatory_questionaire_form_de, Mandatory_questionaire_form_fg, inclusion_form, CT_Scan_Nodule_Form_1, CT_Scan_Nodule_Form_2, annual_study_update_part_a_form, annual_study_update_part_b_form, annual_study_update_medications_form, ProtocolDeviationsForm
from .models import Participant, Breath_Collection, Blood_Collection, lab_processing, Indirect_costs, Stai_y1_2, Veterans_rand_12, Mandatory_questionaire, ct_scan_baseline, ct_scan_t1, ct_scan_t2, ct_scan_t3, ct_scan_t4, ct_scan_t5, ct_scan_t6, ct_scan_t7, ct_scan_nodule_1, ct_scan_nodule_2, ct_scan_nodule_3, ct_scan_nodule_4, ct_scan_nodule_5, Exposure, Exposure2, Exposure3, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg, inclusion, Exposure, Exposure2, Exposure3, Mandatory_questionaire, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg, inclusion, PLCO_score, annual_study_update_part_a, annual_study_update_part_b, annual_study_update_medications, annual_study_update, last_mandatory_questionnaire, Protocol_Deviations, Clinical_Procedures
from django.shortcuts import render, get_object_or_404
from .models import ct_scan_baseline, ct_scan_t1, ct_scan_t2, ct_scan_t3, ct_scan_t4, ct_scan_t5, ct_scan_t6
from django.forms import modelformset_factory
from .forms import biological_relatives_with_cancer_form, plco_score_form, ParticipantSearchForm, CT_Scan_Nodule_Form_3, CT_Scan_Nodule_Form_4, CT_Scan_Nodule_Form_5, annual_study_update_part_b_form, annual_study_update_part_a_form, annual_study_update_form, last_mandatory_questionnaire_form, clinical_procedures_form
##from django.core.exceptions import ObjectDoesNotExist
import csv
from django.http import HttpResponseRedirect as redirect
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.http import JsonResponse
from django.contrib.sessions.models import Session
import re
from django.http import FileResponse
from pdfrw import PdfReader, PdfWriter, IndirectPdfDict
import json
from django.db.models import Q
from django.db import IntegrityError
from django.contrib import messages
from .models import History, biological_relatives_with_cancer
from django.views.generic.edit import CreateView, UpdateView
import logging
from django.contrib.auth.decorators import login_required
from django.contrib.sessions.models import Session
from django.template.loader import get_template
from django.views.generic import View
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
import json
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib import styles
from django.contrib.staticfiles import finders
from django.http import JsonResponse
from xhtml2pdf import pisa
from django.template.loader import get_template
from django.contrib import messages
from django.forms import formset_factory
from django.db import connection
from django.core.exceptions import ObjectDoesNotExist
from django.core.exceptions import FieldError
from django.http import JsonResponse
from django.contrib.auth.models import User
from django.apps import apps
from accounts.models import UserProfile
import os
from django.http import HttpResponse
from django.conf import settings
from docx import Document
from io import BytesIO
from docx2pdf import convert as docx2pdf_convert
from django.contrib.contenttypes.models import ContentType

from django.shortcuts import render
from django.http import HttpResponse
from django.template.loader import render_to_string
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas





def check_participant_number(request):
    if request.method == 'GET':
        participant_number = request.GET.get('participant_number')
        if participant_number:
            try:
                participant = Participant.objects.get(participant_number=participant_number)
                return JsonResponse({'exists': True})
            except Participant.DoesNotExist:
                return JsonResponse({'exists': False})
    return JsonResponse({'error': 'Invalid request'})


def download_data(request):
    
    all_data = []
    all_data.extend(list(inclusion.objects.all()))
    all_data.extend(list(Exposure2.objects.all()))
    all_data.extend(list(Exposure.objects.all()))
    all_data.extend(list(Blood_Collection.objects.all()))
    all_data.extend(list(Breath_Collection.objects.all()))
    all_data.extend(list(Exposure3.objects.all()))
    all_data.extend(list(Indirect_costs.objects.all()))
    all_data.extend(list(Veterans_rand_12.objects.all()))
    all_data.extend(list(Stai_y1_2.objects.all()))
    all_data.extend(list(Mandatory_questionaire.objects.all()))
    all_data.extend(list(Mandatory_questionaire_c.objects.all()))
    all_data.extend(list(Mandatory_questionaire_de.objects.all()))
    all_data.extend(list(Mandatory_questionaire_fg.objects.all()))
    all_data.extend(list(last_mandatory_questionnaire.objects.all()))
    all_data.extend(list(ct_scan_t1.objects.all()))
    all_data.extend(list(ct_scan_t2.objects.all()))
    all_data.extend(list(ct_scan_t3.objects.all()))
    all_data.extend(list(ct_scan_t4.objects.all()))
    all_data.extend(list(ct_scan_t5.objects.all()))
    all_data.extend(list(ct_scan_t6.objects.all()))
    all_data.extend(list(ct_scan_t7.objects.all()))
    all_data.extend(list(ct_scan_nodule_1.objects.all()))
    all_data.extend(list(ct_scan_nodule_2.objects.all()))
    all_data.extend(list(ct_scan_nodule_3.objects.all()))
    all_data.extend(list(ct_scan_nodule_4.objects.all()))
    all_data.extend(list(ct_scan_nodule_5.objects.all()))
    all_data.extend(list(PLCO_score.objects.all()))
    all_data.extend(list(annual_study_update_part_a.objects.all()))
    all_data.extend(list(annual_study_update_medications.objects.all()))
    all_data.extend(list(annual_study_update_part_b.objects.all()))
    all_data.extend(list(annual_study_update.objects.all()))

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="participant_data.csv"'

    writer = csv.writer(response)

    header_row = []
    header_row.extend(inclusion._meta.fields)
    header_row.extend(Exposure2._meta.fields)
    header_row.extend(Exposure._meta.fields)
    header_row.extend(Blood_Collection._meta.fields)
    header_row.extend(Breath_Collection._meta.fields)
    header_row.extend(Exposure3._meta.fields)
    header_row.extend(Indirect_costs._meta.fields)
    header_row.extend(Veterans_rand_12._meta.fields)
    header_row.extend(Stai_y1_2._meta.fields)
    header_row.extend(Mandatory_questionaire._meta.fields)
    header_row.extend(Mandatory_questionaire_c._meta.fields)
    header_row.extend(Mandatory_questionaire_de._meta.fields)
    header_row.extend(Mandatory_questionaire_fg._meta.fields)
    header_row.extend(last_mandatory_questionnaire._meta.fields)
    header_row.extend(ct_scan_t1._meta.fields)
    header_row.extend(ct_scan_t2._meta.fields)
    header_row.extend(ct_scan_t3._meta.fields)
    header_row.extend(ct_scan_t4._meta.fields)
    header_row.extend(ct_scan_t5._meta.fields)
    header_row.extend(ct_scan_t6._meta.fields)
    header_row.extend(ct_scan_t7._meta.fields)
    header_row.extend(ct_scan_nodule_1._meta.fields)
    header_row.extend(ct_scan_nodule_2._meta.fields)
    header_row.extend(ct_scan_nodule_3._meta.fields)
    header_row.extend(ct_scan_nodule_4._meta.fields)
    header_row.extend(ct_scan_nodule_5._meta.fields)
    header_row.extend(PLCO_score._meta.fields)
    header_row.extend(annual_study_update_part_a._meta.fields)
    header_row.extend(annual_study_update_medications._meta.fields)
    header_row.extend(annual_study_update_part_b._meta.fields)
    header_row.extend(annual_study_update._meta.fields)



    writer.writerow(header_row)

    # header_row = []
    # for model in [inclusion, Exposure2, Exposure, Blood_Collection, Breath_Collection, Exposure3, Indirect_costs, Veterans_rand_12, Stai_y1_2, Mandatory_questionaire, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg, last_mandatory_questionnaire, ct_scan_t1, ct_scan_t2, ct_scan_t3, ct_scan_t4, ct_scan_t5, ct_scan_t6, ct_scan_t7, ct_scan_nodule_1, ct_scan_nodule_2, ct_scan_nodule_3, ct_scan_nodule_4, ct_scan_nodule_5, PLCO_score, annual_study_update_part_a, annual_study_update_medications, annual_study_update_part_b, annual_study_update]:
    #     header_row.extend([field.verbose_name for field in model._meta.fields])

    # writer.writerow(header_row)


    # for item in all_data:
    #     row = []
    #     for field in header_row:
    #         try:
    #             value = getattr(item, field.name)
    #             row.append(value)
    #         except AttributeError:
    #             row.append(None)
    #     writer.writerow(row)
    for item in all_data:
        row = []
        for field in header_row:
            try:
                value = getattr(item, field.name)
                if value is None or value == '':
                    value = 'N/A'  # Replace None or empty string with 'N/A'
                row.append(value)
            except AttributeError:
                row.append('N/A')  # Also replace missing attributes with 'N/A'
        writer.writerow(row)

    return response



def download_data_old2(request):

    #models = apps.get_models()

    data_model_1 = inclusion.objects.all().values_list('participant_num', flat=True)
    data_model_2 = Exposure2.objects.all().values_list('participant_num', flat=True)
    data_model_3 = Exposure.objects.all().values_list('participant_num', flat=True)
    data_model_4 = Blood_Collection.objects.all().values_list('participant_num', flat=True)
    data_model_5 = Breath_Collection.objects.all().values_list('participant_num', flat=True)
    data_model_6 = Exposure3.objects.all().values_list('participant_num', flat=True)
    data_model_7 = Indirect_costs.objects.all().values_list('participant_num', flat=True)
    data_model_8 = Veterans_rand_12.objects.all().values_list('participant_num', flat=True)
    data_model_9 = Stai_y1_2.objects.all().values_list('participant_num', flat=True)
    data_model_10 = Mandatory_questionaire.objects.all().values_list('participant_num', flat=True)
    data_model_11 = Mandatory_questionaire_c.objects.all().values_list('participant_num', flat=True)
    data_model_12 = Mandatory_questionaire_de.objects.all().values_list('participant_num', flat=True)
    data_model_13 = Mandatory_questionaire_fg.objects.all().values_list('participant_num', flat=True)
    data_model_14 = last_mandatory_questionnaire.objects.all().values_list('participant_num', flat=True)
    data_model_15 = ct_scan_t1.objects.all().values_list('participant_num', flat=True)
    data_model_16 = ct_scan_t2.objects.all().values_list('participant_num', flat=True)
    data_model_17 = ct_scan_t3.objects.all().values_list('participant_num', flat=True)
    data_model_18 = ct_scan_t4.objects.all().values_list('participant_num', flat=True)
    data_model_19 = ct_scan_t5.objects.all().values_list('participant_num', flat=True)
    data_model_20 = ct_scan_t6.objects.all().values_list('participant_num', flat=True)
    data_model_21 = ct_scan_t7.objects.all().values_list('participant_num', flat=True)
    data_model_22 = ct_scan_nodule_1.objects.all().values_list('participant_num', flat=True)
    data_model_23 = ct_scan_nodule_2.objects.all().values_list('participant_num', flat=True)
    data_model_24 = ct_scan_nodule_3.objects.all().values_list('participant_num', flat=True)
    data_model_25 = ct_scan_nodule_4.objects.all().values_list('participant_num', flat=True)
    data_model_26 = ct_scan_nodule_5.objects.all().values_list('participant_num', flat=True)
    data_model_27 = PLCO_score.objects.all().values_list('participant_num', flat=True)
    data_model_28 = annual_study_update_part_a.objects.all().values_list('participant_num', flat=True)
    data_model_29 = annual_study_update_medications.objects.all().values_list('participant_num', flat=True)
    data_model_30 = annual_study_update_part_b.objects.all().values_list('participant_num', flat=True)
    data_model_31 = annual_study_update.objects.all().values_list('participant_num', flat=True)

    all_data =  [data_model_1      ,    data_model_2      ,    data_model_3      ,    data_model_4      ,    data_model_5      ,    data_model_6      ,    data_model_7      ,    data_model_8      ,    data_model_9      ,    data_model_10      ,    data_model_11      ,    data_model_12      ,    data_model_13      ,     data_model_14      ,    data_model_15      ,    data_model_16      ,    data_model_17      ,    data_model_18      ,    data_model_19      ,    data_model_20      ,    data_model_21      ,    data_model_22      ,    data_model_23      ,    data_model_24      ,    data_model_25      ,    data_model_26      ,    data_model_27      ,    data_model_28      ,    data_model_29      ,    data_model_30      ,    data_model_31   ]

    #all_data = list(data_model_1) + list(data_model_2)+ list(data_model_3)+ list(data_model_4)+ list(data_model_5)+ list(data_model_6)+ list(data_model_7)+ list(data_model_8)+ list(data_model_9)+ list(data_model_10)+ list(data_model_11)+ list(data_model_12)+ list(data_model_13)+  list(data_model_14)+ list(data_model_15)+ list(data_model_16)+ list(data_model_17)+ list(data_model_18)+ list(data_model_19)+ list(data_model_20)+ list(data_model_21)+ list(data_model_22)+ list(data_model_23)+ list(data_model_24)+ list(data_model_25)+ list(data_model_26)+ list(data_model_27)+ list(data_model_28)+ list(data_model_29)+ list(data_model_30)+ list(data_model_31)


    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="participant_data.csv"'


    writer = csv.writer(response)


    header_row = []


    for model in models:
        header_row.extend([field.name for field in model._meta.fields])

    writer.writerow(header_row)


    for item in all_data:
        row = []
        for model in models:

            for field in model._meta.fields:
                try:
                    row.append(getattr(item, field.name))
                except AttributeError:
                    row.append(None)  
        writer.writerow(row)

    return response




def download_data_old (request):
    data_model_1 = inclusion.objects.all().values_list('participant_num', flat=True)
    data_model_2 = Exposure2.objects.all().values_list('participant_num', flat=True)
    data_model_3 = Exposure.objects.all().values_list('participant_num', flat=True)
    data_model_4 = Blood_Collection.objects.all().values_list('participant_num', flat=True)
    data_model_5 = Breath_Collection.objects.all().values_list('participant_num', flat=True)
    data_model_6 = Exposure3.objects.all().values_list('participant_num', flat=True)
    data_model_7 = Indirect_costs.objects.all().values_list('participant_num', flat=True)
    data_model_8 = Veterans_rand_12.objects.all().values_list('participant_num', flat=True)
    data_model_9 = Stai_y1_2.objects.all().values_list('participant_num', flat=True)
    data_model_10 = Mandatory_questionaire.objects.all().values_list('participant_num', flat=True)
    data_model_11 = Mandatory_questionaire_c.objects.all().values_list('participant_num', flat=True)
    data_model_12 = Mandatory_questionaire_de.objects.all().values_list('participant_num', flat=True)
    data_model_13 = Mandatory_questionaire_fg.objects.all().values_list('participant_num', flat=True)
    data_model_14 = last_mandatory_questionnaire.objects.all().values_list('participant_num', flat=True)
    data_model_15 = ct_scan_t1.objects.all().values_list('participant_num', flat=True)
    data_model_16 = ct_scan_t2.objects.all().values_list('participant_num', flat=True)
    data_model_17 = ct_scan_t3.objects.all().values_list('participant_num', flat=True)
    data_model_18 = ct_scan_t4.objects.all().values_list('participant_num', flat=True)
    data_model_19 = ct_scan_t5.objects.all().values_list('participant_num', flat=True)
    data_model_20 = ct_scan_t6.objects.all().values_list('participant_num', flat=True)
    data_model_21 = ct_scan_t7.objects.all().values_list('participant_num', flat=True)
    data_model_22 = ct_scan_nodule_1.objects.all().values_list('participant_num', flat=True)
    data_model_23 = ct_scan_nodule_2.objects.all().values_list('participant_num', flat=True)
    data_model_24 = ct_scan_nodule_3.objects.all().values_list('participant_num', flat=True)
    data_model_25 = ct_scan_nodule_4.objects.all().values_list('participant_num', flat=True)
    data_model_26 = ct_scan_nodule_5.objects.all().values_list('participant_num', flat=True)
    data_model_27 = PLCO_score.objects.all().values_list('participant_num', flat=True)
    data_model_28 = annual_study_update_part_a.objects.all().values_list('participant_num', flat=True)
    data_model_29 = annual_study_update_medications.objects.all().values_list('participant_num', flat=True)
    data_model_30 = annual_study_update_part_b.objects.all().values_list('participant_num', flat=True)
    data_model_31 = annual_study_update.objects.all().values_list('participant_num', flat=True)



     
    all_data = list(data_model_1) + list(data_model_2)+ list(data_model_3)+ list(data_model_4)+ list(data_model_5)+ list(data_model_6)+ list(data_model_7)+ list(data_model_8)+ list(data_model_9)+ list(data_model_10)+ list(data_model_11)+ list(data_model_12)+ list(data_model_13)+  list(data_model_14)+ list(data_model_15)+ list(data_model_16)+ list(data_model_17)+ list(data_model_18)+ list(data_model_19)+ list(data_model_20)+ list(data_model_21)+ list(data_model_22)+ list(data_model_23)+ list(data_model_24)+ list(data_model_25)+ list(data_model_26)+ list(data_model_27)+ list(data_model_28)+ list(data_model_29)+ list(data_model_30)+ list(data_model_31)

    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="participant_data.csv"'

    writer = csv.writer(response)
    writer.writerow(['Participant Number']) 

    for item in all_data:
        row = [getattr(item, field.name) for field in YourModel1._meta.fields]
        writer.writerow(row)
    return response



def get_current_user(request):
    if request.user.is_authenticated:
        user_profile = UserProfile.objects.get(user=request.user)
        site = user_profile.site
        current_user = {
            'username': request.user.username,
            'email': request.user.email,
            'site': site,
            
        }
        return JsonResponse(current_user)
    else:
        return JsonResponse({'error': 'User not logged in'}, status=401)






def download_pdf(request):
    pdf_path = finders.find('your_file.pdf')
    print('Working here 1')
    if pdf_path:
        with open(pdf_path, 'rb') as pdf_file:
            print('working here')
            response = HttpResponse(pdf_file.read(), content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="your_file.pdf"'
            return response
    else:
        return HttpResponse("PDF file not found.", status=404)



    

class GeneratePDF2(View):
    def get(self, request, *args, **kwargs):
        form_data = get_form_data(request)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="data.pdf"'

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        elements = []
        elements.append(Paragraph("Form Data PDF", styles['Title']))
        elements.append(Paragraph("Below is the data submitted in the form:", styles['Normal']))
        
        for field, value in form_data.items():
            field_line = f"<strong>{field}:</strong> {value}<br/>"
            elements.append(Paragraph(field_line, styles['Normal']))

        doc.build(elements)
        pdf = buffer.getvalue()
        buffer.close()

        response.write(pdf)
        return response


class GeneratePDF(View):
    def get(self, request, *args, **kwargs):
        form_data_json = request.session.get('form_data')
        if form_data_json:
            form_data = json.loads(form_data_json)
            form_data_string = "\n".join([f"{key}: {value}" for key, value in form_data.items()])

            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="data.pdf"'

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()

            elements = []
            elements.append(Paragraph("Form Data PDF", styles['Title']))
            elements.append(Paragraph("Below is the data submitted in the form:", styles['Normal']))
            elements.append(Paragraph(form_data_string, styles['Normal'])) 

            doc.build(elements)
            pdf = buffer.getvalue()
            buffer.close()

            response.write(pdf)
            return response
        else:
            return HttpResponse("Form data not available.")


def generate_pdf_oldNotWorking(data):
    print('working here')
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    for field_name, value in data.items():
        if field_name != 'csrfmiddlewaretoken':
            field_label = field_name.replace('_', ' ').title()
            elements.append(f'{field_label}: {value}')

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_pdf(data):
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    paragraph_style = styles.getSampleStyleSheet()['Normal']
    for field_name, value in data.items():
        if field_name != 'csrfmiddlewaretoken':
            field_label = field_name.replace('_', ' ').title()
            paragraph_text = f'<b>{field_label}:</b> {value}'
            paragraph = Paragraph(paragraph_text, style=paragraph_style)
            elements.append(paragraph)

    doc.build(elements)
    buffer.seek(0)
    print('working pdf here')
    return buffer

def get_form_data(request):
    form_data = {
        'Visit Date': request.POST.get('{form.visit_date}', ''),
        'Postal Code': request.POST.get('postal_code', ''),
        'Date of Birth': request.POST.get('date_of_birth', ''),
        'Current Age': request.POST.get('current_age', ''),
        'Current Height': request.POST.get('current_height', ''),
        'Current Height Unit': request.POST.get('current_height_unit', ''),
        'Current Weight': request.POST.get('current_weight', ''),
        'Current Weight Unit': request.POST.get('current_weight_unit', ''),
        'Sex at Birth': request.POST.get('sex_birth', ''),
        'Gender Identity': request.POST.get('gender_identity', ''),
        'Ethnicity': request.POST.get('ethnicity', ''),
        'Ethnicity (Other)': request.POST.get('ethnicity_other', ''),
        'Born in Canada': request.POST.get('born_in_canada', ''),
        'Year Moved to Canada': request.POST.get('year_moved_to_canada', ''),
        'Birthplace': request.POST.get('birthplace', ''),
        'Highest Education Level': request.POST.get('highest_education_lvl', ''),
        'Highest Education Level (Other)': request.POST.get('highest_education_lvl_other', ''),
    }
    return form_data

@login_required(login_url='login')
def index(request):
    form = ParticipantForm
    return render(request, 'DataEntry/index.html', {'form': form})

@login_required(login_url='login')
def index_participant(request, participant_id=None):
    # Get all models in the project
    models = apps.get_models()

    # Initialize a defaultdict to store timestamps for each model
    last_entries = {}

    # Create a cursor to execute raw SQL queries
    cursor = connection.cursor()

    # Iterate through each model to check its historical table
    for model in models:
        # Exclude system tables
        if model._meta.app_label != 'django' and not model._meta.model_name.startswith('auth_'):
            # Get the historical table name
            historical_table_name = f'{model._meta.app_label.lower()}_{model._meta.model_name}'

            # Check if the historical table exists in the database
            cursor.execute(f"SHOW TABLES LIKE '{historical_table_name}'")
            table_exists = cursor.fetchone()

            if table_exists:
                # Debug line to check if history_date column exists in the table
                cursor.execute(f"SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = '{historical_table_name}' AND COLUMN_NAME = 'history_date'")
                history_date_column_exists = cursor.fetchone()

                if history_date_column_exists:
                    order_by_clause = "ORDER BY history_date DESC"
                else:
                    order_by_clause = ""

                cursor.execute(f"SELECT COUNT(*) FROM {historical_table_name}")
                record_count = cursor.fetchone()[0]

                if record_count > 0:
                    # Historical table has records, proceed with querying for updated records
                    sql_query = f"SELECT * FROM {historical_table_name} {order_by_clause} LIMIT 1"
                    cursor.execute(sql_query)
                    latest_record = cursor.fetchone()

                    if latest_record:
                        print(f"Most recent update found in {historical_table_name}: {latest_record}")
                        # Add the timestamp to the corresponding model in the dictionary
                        history_date_index = None
                        for index, column_name in enumerate(cursor.description):
                            if column_name[0] == 'history_date':
                                history_date_index = index
                                break

                        if history_date_index is not None:
                            history_date = latest_record[history_date_index]
                            last_entries[model.__name__] = history_date
                        else:
                            pass
                    else:
                        pass
            else:
                pass

    form = ParticipantForm()
    request.session['participant_id'] = participant_id

    search_results = []

    models_list = [ct_scan_baseline, ct_scan_t1, ct_scan_t2, ct_scan_t3, ct_scan_t4, ct_scan_t5, ct_scan_t6,
                   lab_processing, Breath_Collection, Blood_Collection, Exposure, Exposure2, Exposure3, Mandatory_questionaire,
                   Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg, inclusion,
                   PLCO_score, annual_study_update_part_a, annual_study_update_part_b, annual_study_update_medications, annual_study_update,
                   last_mandatory_questionnaire, Protocol_Deviations, Clinical_Procedures, ct_scan_nodule_1, ct_scan_nodule_2, ct_scan_nodule_3,
                   ct_scan_nodule_4, ct_scan_nodule_5, Indirect_costs]

    for model_class in models_list:
        try:
            if 'participant_num' in [f.name for f in model_class._meta.get_fields()] and hasattr(model_class, 'history'):
                entries = model_class.objects.filter(participant_num=participant_id)
                for entry in entries:
                    search_results.append({'model_name': model_class.__name__, 'entry': entry})
                    content_type = ContentType.objects.get_for_model(model_class)
                    history = model_class.history.filter(instance_id=entry.id, content_type=content_type)
                    if history.exists():
                        print('History:')
                        for record in history:
                            timestamps = [record.history_date for record in history]
                            last_entries[model_class.__name__] = timestamps
                    else:
                        pass
            else:
                pass
        except (FieldError, model_class.DoesNotExist):
            pass
    formatted_entries = {}
    for model_name, timestamp in last_entries.items():
        formatted_model_name = model_name.replace('Historical', '').replace('_', ' ').capitalize()
        formatted_entries[formatted_model_name] = timestamp

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials
    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
        'last_entries': last_entries,
        'search_results': search_results,
        'formatted_entries': formatted_entries,
    }

    return render(request, 'DataEntry/index_participant.html', context)

from django.shortcuts import render, redirect
from .models import Participant, UploadedPDF
from .forms import UploadPDFForm

def add_mandatory_questionaire_dashboard_old(request, participant_id):
    participant = Participant.objects.get(participant_number=participant_id)
    if request.method == 'POST':
        form = UploadPDFForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_pdf = form.save(commit=False)
            uploaded_pdf.participant = participant
            uploaded_pdf.save()
            return redirect('list_pdf', participant_id=participant_id)
    else:
        form = UploadPDFForm()
    uploaded_pdfs = UploadedPDF.objects.filter(participant=participant)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    participant = participant_id
    print(participant)
    inclusion_data = inclusion.objects.filter(participant_num=participant_id).exists()
    questionnaire_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).exists()
    questionnaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant_id).exists()
    questionnaire_de_data = last_mandatory_questionnaire.objects.filter(participant_num=participant_id).exists()
    exposure_data = Exposure2.objects.filter(participant_num=participant_id).exists()
    print("Questionnaire Data:", questionnaire_data)
    print("Questionnaire C Data:", questionnaire_c_data)
    print("Questionnaire D Data:", questionnaire_de_data)
    print("Exposure Data:", exposure_data)

    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
        'participant': participant,
        'inclusion_data': inclusion_data,
        'questionnaire_data': questionnaire_data,
        'questionnaire_c_data': questionnaire_c_data,
        'questionnaire_d_data': questionnaire_de_data,
        'exposure_data': exposure_data,
        'uploaded_pdfs': uploaded_pdfs,
    }

    return render(request, 'DataEntry/mandatory_questionaire_dashboard.html', context)
    
from .models import mandatory_questionaire_dashboard
from .forms import mandatory_questionaire_dashboard_form

@login_required
def add_mandatory_questionaire_dashboard(request, participant_id=None):
    participant = get_object_or_404(Participant, pk=participant_id)
    
    if request.method == 'POST':
        form = mandatory_questionaire_dashboard_form(request.POST, request.FILES)
        if form.is_valid():
            try:
                mandatory_questionaire_dashboard_obj = form.save(commit=False)
                mandatory_questionaire_dashboard_obj.participant_num = participant
                mandatory_questionaire_dashboard_obj.save()
                return redirect('mandatory_questionaire_dashboard', participant_id=participant_id)
            except Exception as e:
                print("Error saving file:", e)
                traceback.print_exc()
        else:
            print("Form errors:", form.errors)
    else:
        form = mandatory_questionaire_dashboard_form()

    uploaded_files = mandatory_questionaire_dashboard.objects.filter(participant_num=participant)
    inclusion_data = inclusion.objects.filter(participant_num=participant_id).exists()
    questionnaire_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).exists()
    questionnaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant_id).exists()
    questionnaire_de_data = last_mandatory_questionnaire.objects.filter(participant_num=participant_id).exists()
    exposure_data = Exposure2.objects.filter(participant_num=participant_id).exists()
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant).first()
    date_of_birth = base_data.date_of_birth if base_data else 'N/A'
    initials = base_data.initials if base_data else 'N/A'

    context = {
        'form': form,
        'uploaded_files': uploaded_files,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
        'inclusion_data': inclusion_data,
        'questionnaire_data': questionnaire_data,
        'questionnaire_c_data': questionnaire_c_data,
        'questionnaire_de_data': questionnaire_de_data,
        'exposure_data': exposure_data,
        
    }

    return render(request, 'DataEntry/mandatory_questionaire_dashboard.html', context)

def delete_mandatory_dashboard_file(request, file_id):
    if request.method == 'POST':
        file_to_delete = get_object_or_404(mandatory_questionaire_dashboard, pk=file_id)
        participant_id = file_to_delete.participant_num_id  # Retrieve participant ID before deleting
        file_to_delete.delete()
    return redirect('mandatory_questionaire_dashboard', participant_id=participant_id)



from django.shortcuts import get_object_or_404

def delete_pdf(request, pdf_id):
    uploaded_pdf = get_object_or_404(UploadedPDF, id=pdf_id)
    print("PDF ID:", pdf_id)
    print("Uploaded PDF:", uploaded_pdf)
    participant_id = uploaded_pdf.participant.participant_number
    print("Participant ID:", participant_id)
    uploaded_pdf.delete()

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    participant = participant_id
    print(participant)
    inclusion_data = inclusion.objects.filter(participant_num=participant_id).exists()
    questionnaire_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).exists()
    questionnaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant_id).exists()
    questionnaire_de_data = last_mandatory_questionnaire.objects.filter(participant_num=participant_id).exists()
    exposure_data = Exposure2.objects.filter(participant_num=participant_id).exists()
    print("Questionnaire Data:", questionnaire_data)
    print("Questionnaire C Data:", questionnaire_c_data)
    print("Questionnaire D Data:", questionnaire_de_data)
    print("Exposure Data:", exposure_data)

    context = {

        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
        'participant': participant,
        'inclusion_data': inclusion_data,
        'questionnaire_data': questionnaire_data,
        'questionnaire_c_data': questionnaire_c_data,
        'questionnaire_d_data': questionnaire_de_data,
        'exposure_data': exposure_data,
    }

    return redirect('DataEntry/list_pdf', context)

def list_pdf(request, participant_id):
    participant = Participant.objects.get(participant_number=participant_id)
    uploaded_pdfs = UploadedPDF.objects.filter(participant=participant)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    participant = participant_id
    print(participant)
    inclusion_data = inclusion.objects.filter(participant_num=participant_id).exists()
    questionnaire_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).exists()
    questionnaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant_id).exists()
    questionnaire_de_data = last_mandatory_questionnaire.objects.filter(participant_num=participant_id).exists()
    exposure_data = Exposure2.objects.filter(participant_num=participant_id).exists()
    print("Questionnaire Data:", questionnaire_data)
    print("Questionnaire C Data:", questionnaire_c_data)
    print("Questionnaire D Data:", questionnaire_de_data)
    print("Exposure Data:", exposure_data)

    context = {
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
        'participant': participant,
        'inclusion_data': inclusion_data,
        'questionnaire_data': questionnaire_data,
        'questionnaire_c_data': questionnaire_c_data,
        'questionnaire_d_data': questionnaire_de_data,
        'exposure_data': exposure_data,
        'uploaded_pdfs': uploaded_pdfs,
    }
    return render(request, 'DataEntry/list_pdf.html', context)
# @login_required(login_url='login')
# def add_mandatory_questionaire_dashboard(request, participant_id=None):
#     form = pdf_uploads_questionnaire_form(request.POST, request.FILES)


#     request.session['participant_id'] = participant_id
    
#     if participant_id:
#         request.session['participant_id'] = participant_id
#     if request.method == 'POST':
#         uploaded_file = request.FILES['document']
#         fs = FileSystemStorage()
#         name = fs.save(uploaded_file.name, uploaded_file)
#         context['url'] = fs.url(name)

#     participant = participant_id
#     print(participant)
#     inclusion_data = inclusion.objects.filter(participant_num=participant_id).exists()
#     questionnaire_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).exists()
#     questionnaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant_id).exists()
#     questionnaire_de_data = last_mandatory_questionnaire.objects.filter(participant_num=participant_id).exists()
#     exposure_data = Exposure2.objects.filter(participant_num=participant_id).exists()


#     base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
#     if base_data:
#         date_of_birth = base_data.date_of_birth
#         initials = base_data.initials

#     else:
#         date_of_birth = 'N/A'
#         initials = 'N/A'

#     print("Questionnaire Data:", questionnaire_data)
#     print("Questionnaire C Data:", questionnaire_c_data)
#     print("Questionnaire D Data:", questionnaire_de_data)
#     print("Exposure Data:", exposure_data)

#     context = {
#         'form': form,
#         'participant_id': participant_id,
#         'initials': initials,
#         'date_of_birth': date_of_birth,
#         'participant': participant,
#         'inclusion_data': inclusion_data,
#         'questionnaire_data': questionnaire_data,
#         'questionnaire_c_data': questionnaire_c_data,
#         'questionnaire_d_data': questionnaire_de_data,
#         'exposure_data': exposure_data,
#     }

#     return render(request, 'DataEntry/mandatory_questionaire_dashboard.html', context)


from reportlab.lib.pagesizes import letter
from django.shortcuts import render
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML




from django.views.generic import TemplateView


def download_mandatory_questionaire_pdf(request, participant_id):
    # Fetch data from models based on participant_id
    form = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    form1 = Mandatory_questionaire_c.objects.filter(participant_num=participant_id).first()
    form2 = last_mandatory_questionnaire.objects.filter(participant_num=participant_id).first()
    form3 = biological_relatives_with_cancer.objects.filter(participant_num=participant_id).first()
    
    # Render HTML template with data
    context = {
        'form': form,
        'form1': form1,
        'form2': form2,
        'form3': form3,
    }
    template = render(request, 'DataEntry/full_mandatory_questionnaire_pdf.html', context)

    # Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{participant_id}_Mandatory_Questionnaire.pdf"'

    # Generate PDF from HTML template
    pisa_status = pisa.CreatePDF(template.content, dest=response)
    if pisa_status.err:
        return HttpResponse('Error generating PDF')

    return response




def download_mandatory_questionaire_pdf_old2(request, participant_id):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="form_data.pdf"'
    participant_num = participant_id

    # Handle form submission and extract data
    form = Mandatory_questionaire.objects.filter(participant_num=participant_num).first()
    form1 = Mandatory_questionaire_c.objects.filter(participant_num=participant_num).first()
    form2 = Mandatory_questionaire_de.objects.filter(participant_num=participant_num).first()
 
    # Render HTML template with submitted data
    html_string = render_to_string('DataEntry/full_mandatory_questionnaire_pdf.html', {'form': form, 'participant_id': participant_num})
    html_string2 = render_to_string('DataEntry/full_mandatory_questionnaire_pdf.html', {'form1': form1, 'participant_id': participant_num})
    html_string3 = render_to_string('DataEntry/full_mandatory_questionnaire_pdf.html', {'form2': form2, 'participant_id': participant_num})
    # Generate PDF from HTML

    final_html = html_string + html_string2 + html_string3
    pdf = HTML(string=final_html).write_pdf()
    #pdf = HTML(string=html_string).write_pdf()

    filename = f'Mandatory_Questionnaire_{participant_id}.pdf'
    # Return PDF as response
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response



import base64

def download_optional_questionaire_pdf(request, participant_id):
    # image_path = os.path.join(settings.BASE_DIR, 'ideal-logo-blue-white.png')
    # with open(image_path, 'rb') as f:
    #     image_data = f.read()
    #     image_base64 = base64.b64encode(image_data).decode()
    # Fetch data from models based on participant_id
    form2 = Exposure2.objects.filter(participant_num=participant_id).first()
    form3 = Exposure3.objects.filter(participant_num=participant_id).first()
    
    # Render HTML template with data
    context = {
        'form2': form2,
        'form3': form3,
        

    }
    template = render(request, 'DataEntry/exposure_form_pdf.html', context)

    # Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{participant_id}_Optional_Questionnaire.pdf"'

    # Generate PDF from HTML template
    pisa_status = pisa.CreatePDF(template.content, dest=response)
    if pisa_status.err:
        return HttpResponse('Error generating PDF')

    return response






def download_optional_questionaire_pdf_old(request, participant_id):
    participant_id = participant_id
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="form_data.pdf"'
    participant_num = participant_id

    # Handle form submission and extract data
    form = Exposure2.objects.filter(participant_num=participant_num).first()
    form2 = Exposure3.objects.filter(participant_num=participant_num).first()

 
    # Render HTML template with submitted data
    html_string = render_to_string('DataEntry/exposure_form.html', {'form': form, 'participant_id': participant_num})
    html_string2 = render_to_string('DataEntry/residence_history_form.html', {'form': form2, 'participant_id': participant_num})

    # Generate PDF from HTML

    final_html = html_string + html_string2 
    pdf = HTML(string=final_html).write_pdf()
    #pdf = HTML(string=html_string).write_pdf()
    filename = f'Optional_Questionnaire_{participant_id}.pdf'
    # Return PDF as response
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response



    # return render(request, 'form_template.html')



    # # Get all models and their fields dynamically
    # participant_id = request.GET.get('participant_id')
    # models = [Mandatory_questionaire, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg]


    # data_to_include = []
    # for model_class in models:
    #     print('working')
    #     # Filter data based on participant_id
    #     model_data = model_class.objects.filter(participant_num=participant_id)
    #     for instance in model_data:
    #         print('working2')
    #         field_data = {}
    #         for field in instance._meta.fields:
    #             field_data[field.name] = getattr(instance, field.name)
    #         data_to_include.append(field_data)

    # # Create PDF document
    # response = HttpResponse(content_type='application/pdf')
    # response['Content-Disposition'] = 'attachment; filename="Mandatory Questionaire.pdf"'

    # pdf_canvas = canvas.Canvas(response, pagesize=letter)
    # print("Data to include:", data_to_include)
    # # Add data to PDF
    # y_position = 750  # Initial y-position for text
    # for item in data_to_include:
    #     for field_name, field_value in item.items():
    #         pdf_canvas.drawString(100, y_position, f'{field_name}: {field_value}')
    #         y_position -= 20  # Adjust y-position for next field
    #     y_position -= 20  # Add space between model instances

    # # Save PDF
    # pdf_canvas.save()
    # return response



    # if 'generate_pdf' in request.GET:
    #     participant_id = request.GET.get('participant_id')
    #     models = [Mandatory_questionaire, Mandatory_questionaire_c, Mandatory_questionaire_de, Mandatory_questionaire_fg]

    #     data_to_include = []
    #     for model_class in models:
            
    #         model_data = model_class.objects.filter(participant_num=participant_id)
    #         for instance in model_data:
    #             field_data = {}
    #             for field in instance._meta.fields:
    #                 field_data[field.name] = getattr(instance, field.name)
    #             data_to_include.append(field_data)

        
    #     response = HttpResponse(content_type='application/pdf')
    #     response['Content-Disposition'] = 'attachment; filename="Mandatory Questionaire.pdf"'

    #     pdf_canvas = canvas.Canvas(response, pagesize=letter)

        
    #     y_position = 750  
    #     for item in data_to_include:
    #         for field_name, field_value in item.items():
    #             pdf_canvas.drawString(100, y_position, f'{field_name}: {field_value}')
    #             y_position -= 20 
    #         y_position -= 20  

    #     # Save PDF
    #     pdf_canvas.save()
    #     return response




@login_required(login_url='login')
def search_and_add(request):
    form = ParticipantSearchForm()

    if request.method == 'POST':
        form = ParticipantSearchForm(request.POST)
        if form.is_valid():
            participant_number = form.cleaned_data['participant_number']

            try:
                participant = Participant.objects.get(participant_number=participant_number)
                request.session['participant_num'] = participant.participant_number
                return render(request, 'search_result.html', {'participant': participant})
            except Participant.DoesNotExist:
                pass
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    return render(request, 'search_page.html', context)

@login_required(login_url='login')
def add_inclusion(request):
    participant_form = ParticipantForm(request.POST)
    if request.method == 'POST':
        if participant_form.is_valid():
            print('valid')
            participant_num = participant_form.cleaned_data['participant_number']
            participant_num = participant_num.replace('-', '0')
            participant, created = Participant.objects.get_or_create(participant_number=participant_num)
            if created:
                participant_num = participant_form.cleaned_data['participant_number']
                participant_num = participant_num.replace('-', '0')
                participant, created = Participant.objects.get_or_create(participant_number=participant_num)
                return HttpResponseRedirect(reverse('add_inclusion_participant', args=[participant_num]))

        else:
            participant_num = request.POST.get('participant_number')
            return HttpResponseRedirect(reverse('add_inclusion_participant', args=[participant_num]))
        
    else:
        print(participant_form.errors)
        participant_form = ParticipantForm()

    context = {
        'participant_form': participant_form,

    }

    return render(request, 'DataEntry/inclusion.html', context)


def generate_pdf(request, participant_id):
    print("Generating PDF for participant ID:", participant_id)
    participant = get_object_or_404(Participant, participant_number=participant_id)
    inclusion_data = inclusion.objects.filter(participant_num=participant).first()

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{participant_id}_report.pdf"'

    p = canvas.Canvas(response)

    p.drawString(100, 750, "Participant Report")
    p.drawString(100, 730, f"Participant ID: {participant_id}")
    p.drawString(100, 710, "Participant Data:")

    if inclusion_data:
        p.drawString(100, 690, f"Field 1: {inclusion_data.field1}")
        p.drawString(100, 670, f"Field 2: {inclusion_data.field2}")
    p.showPage()
    p.save()
    return response


@login_required(login_url='login')
def add_inclusion_participant(request, participant_id=None):
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        inclusion_data = inclusion.objects.filter(participant_num=participant).first()
        print(participant)
        print('the participant number is: ')
        print(participant_id)

    if request.method == 'POST':
            form = inclusion_form(request.POST, request.FILES, instance=inclusion_data)
            form.participant_num = participant_id
            if form.is_valid():
                print('form is valid')
                print(request.FILES)
                if 'consent_form' in form.fields:
                    consent_form_url = form.cleaned_data['consent_form']
                    form.instance.consent_form_path = consent_form_url
                form.save(commit=False)
                form.participant_number = participant
                form.save()
                print('data saved')
                return HttpResponseRedirect(reverse('add_inclusion_participant', args=[participant_id]) + '?submitted=True')
            
    elif 'generate_pdf' in request.POST:
        return generate_pdf(request, participant_id)
    elif 'submit' in request.POST:
        print('something')
    
    elif 'continue' in request.POST:
        print(request.POST)
        form = inclusion_form(request.POST, request.FILES, instance=inclusion_data)
        if form.is_valid():
            # Corrected form.instance assignments

            if 'consent_form' in form.fields:
                consent_form_url = form.cleaned_data['consent_form']
                form.instance.consent_form_path = consent_form_url
            form.instance.participant_number = participant  # Corrected attribute assignment
            form.save()
            print('data saved')
            return HttpResponseRedirect(reverse('add_mandatory_questionaire', args=[participant_id]))

    else:
        form = inclusion_form(instance=inclusion_data)
        print('errors here')
        print(form.errors)
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials
    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,

    }
    print(form.errors)
    return render(request, 'DataEntry/inclusion_participant.html', context)


@login_required(login_url='login')
def add_inclusion_participant_old(request, participant_id=None):
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        inclusion_data = inclusion.objects.filter(participant_num=participant).first()
        print(participant)
        print('the participant number is: ')
        print(participant_id)

    if not participant_id:
        participant = None
        inclusion_data = None
        print('reaching here')
        print(participant_id)

    if request.method == 'POST':
        if 'generate_pdf' in request.POST:
            return generate_pdf(request, participant_id)
        elif 'submit' in request.POST:

            form = inclusion_form(request.POST, request.FILES, instance=inclusion_data)
            form.participant_num = participant_id
            if form.is_valid():
                print('form is valid')
                print(request.FILES)
                if 'consent_form' in form.fields:
                    print('reaching here')
                    consent_form_url = form.cleaned_data['consent_form']
                    form.instance.consent_form_path = consent_form_url
                form.participant_number = participant_id
                form.save()
                return HttpResponseRedirect(reverse('add_inclusion_participant', args=[participant_id]) + '?submitted=True')


    
    else:
        form = inclusion_form(instance=inclusion_data)
        print(form.fields)
        print(form.errors)

    context = {
        'form': form,
    }
    print(form.errors)
    return render(request, 'DataEntry/inclusion_participant.html', context)
    

def add_inclusion_participant_old(request, participant_num=None):
    participant = None 
    form = inclusion_form()
    if request.method == 'POST':
       
        if form.is_valid():
            form.save()
            return HttpResponseRedirect(reverse('add_inclusion_participant', args=[participant_num]) + '?submitted=True')
            
    else:
        
        form = inclusion_form(instance=participant)  
        print(form.errors)

    context = {
        'form': form,
         
    }

    return render(request, 'DataEntry/inclusion_participant.html', context)


def add_inclusion_oldSept29(request, participant_num=None):
    participant_already_exists = False
    
    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)
        
        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            cleaned_participant_num = participant_num.replace('-', '0')
            
            participant, created = Participant.objects.get_or_create(
                participant_number=cleaned_participant_num
                
            )
            if created:
                return redirect(f'/DataEntry/add_inclusion/{cleaned_participant_num}/')
            else:
                participant_already_exists = True
                print("This participant number has already been used")
        else:
            print("Participant Form Errors:", participant_form.errors)
            return JsonResponse({'error': participant_form.errors}, status=400)

    else:
        participant_form = ParticipantForm()

    context = {
        'participant_form': participant_form,
        'participant_already_exists': participant_already_exists,
    }

    return render(request, 'DataEntry/inclusion.html', context)


def add_inclusion_workingParticipantOnly(request):
    if request.method == 'POST':
        participant_form = ParticipantForm(request.POST)

        if participant_form.is_valid():
            participant_num = participant_form.cleaned_data['participant_number']
            print(participant_num)
            cleaned_participant_num = participant_num.replace('-', '0')
            print(cleaned_participant_num)
            participant, created = Participant.objects.get_or_create(
                participant_number=cleaned_participant_num
            )

        else:
            print("Participant Form Errors:", participant_form.errors)
            return JsonResponse({'error': participant_form.errors}, status=400)

    else:
        participant_form = ParticipantForm()

    context = {
        'participant_form': participant_form,
    }
    return render(request, 'DataEntry/inclusion.html', context)




def generate_word_document(data):
    doc = Document()
    doc.add_heading('Mandatory Questionnaire', level=1)
    return doc

@login_required(login_url='login')
def add_mandatory_questionaire(request, participant_id=None):
    # BiologicalRelativesFormSet = modelformset_factory(
    # biological_relatives_with_cancer,
    # form=biological_relatives_with_cancer_form,  
    # extra=1,
    # )
    print("View accessed!")
    if participant_id: 
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_data = Mandatory_questionaire.objects.filter(participant_num=participant).first()
        biological_relatives_with_cancer_data = biological_relatives_with_cancer.objects.filter(participant_num=participant).first()
        print('the participant number is: ', participant_id)
    
    #RelativeCancerFormSet = modelformset_factory(biological_relatives_with_cancer, form=biological_relatives_with_cancer_form, extra=1)
    
    if request.method == 'POST':
        # formset = BiologicalRelativesFormSet(request.POST, prefix='relatives')
        form = Mandatory_questionaire_form(request.POST, instance=mandatory_questionaire_data)
        form2 = biological_relatives_with_cancer_form(request.POST, instance=biological_relatives_with_cancer_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')

        if form.is_valid() and form2.is_valid():
            mandatory_questionaire = form.save(commit=False)
            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            mandatory_questionaire.participant_num = participant
            mandatory_questionaire.save()

            form2.instance.participant_num = participant  
            form2.save()

            # instances = formset.save(commit=False)
            # for instance in instances:
            #     instance.mandatory_questionaire = mandatory_questionaire
            #     instance.save()

            word_doc = generate_word_document(form.cleaned_data)
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.docx'
            word_doc.save(response)
            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_mandatory_questionaire', args=[participant_id]) + '?submitted=True')
        

        
        else:
            print(form.errors)
            print(form2.errors)
            #print(formset.errors)

    elif 'generate_pdf' in request.GET:  
        print('button pressed')
        form = Mandatory_questionaire_form(instance=mandatory_questionaire_data) 
        
        if mandatory_questionaire_data:
            print('yes')
            #formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.filter(mandatory_questionaire=mandatory_questionaire_data), prefix='relatives')
        else:
            print('no')
            #formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.none(), prefix='relatives')
         # Validate the form
        pdf_data = form.instance 
        pdf_buffer = generate_mandatory_questionnaire_pdf(pdf_data, request)
        print(pdf_buffer)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        print(response)
        mandatory_questionaire_obj = Mandatory_questionaire.objects.get(participant_num=participant_id)
        participant_num = mandatory_questionaire_obj.participant_num  
        response['Content-Disposition'] = f'attachment; filename={participant_num}_mandatory_questionnaire.pdf'
        print('should be wokring here')
        return response
    
    elif 'generate_csv' in request.GET:
        form = Mandatory_questionaire_form(instance=mandatory_questionaire_data)
        mandatory_questionaire_obj = Mandatory_questionaire.objects.get(participant_num=participant_id)
        participant_num = mandatory_questionaire_obj.participant_num    
        
        csv_buffer = generate_mandatory_questionnaire_csv(form, request)
        response = HttpResponse(csv_buffer.getvalue(), content_type='application/csv')
        response['Content-Disposition'] = f'attachment; filename={participant_num}_mandatory_questionnaire.csv'
        return response
    
    elif 'save_continue' in request.GET:
        form = Mandatory_questionaire_form(request.POST, instance=mandatory_questionaire_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')
        if form.is_valid():
            mandatory_questionaire = form.save(commit=False)
            mandatory_questionaire.participant_num = participant
            mandatory_questionaire.save()


        return HttpResponseRedirect(reverse('add_mandatory_questionaire_c', args=[participant_id]))

    else:
        print("Getting Data")
        print("Mandatory Questionnaire Data:", mandatory_questionaire_data)
        print("Biological Relatives with Cancer:", biological_relatives_with_cancer.objects.filter(participant_num=participant_id).first())

        form = Mandatory_questionaire_form(instance=mandatory_questionaire_data)
        form2 = biological_relatives_with_cancer_form(instance=biological_relatives_with_cancer_data)
        if mandatory_questionaire_data:
            print('yes')
            #formset = BiologicalRelativesFormSet(queryset=Mandatory_questionaire.objects.filter(participant_num=participant_id), prefix='relatives')
            
        else:
            print('no')
            #formset = BiologicalRelativesFormSet(queryset=Mandatory_questionaire.objects.filter(participant_num=participant_id), prefix='relatives')
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'form2': form2,
        'participant_id': participant_id,
        'mandatory_questionaire_data' : mandatory_questionaire_data,
        'initials': initials,
        'date_of_birth': date_of_birth,

    }

    print(Mandatory_questionaire.date_of_birth)
    print(form.errors)
    #print(formset.errors)
    #print(request.POST)
    return render(request, 'DataEntry/mandatory_questionaire.html', context)


@login_required(login_url='login')
def add_indirect_costs(request, participant_id=None):
    print("Indirect Costs View accessed!")
    if participant_id: 
        participant = get_object_or_404(Participant, participant_number=participant_id)
        indirect_costs_data = Indirect_costs.objects.filter(participant_num=participant).first()
        print('the participant number is: ', participant_id)
    
    #RelativeCancerFormSet = modelformset_factory(biological_relatives_with_cancer, form=biological_relatives_with_cancer_form, extra=1)
    
    if request.method == 'POST':
        form = Indirect_costs_form(request.POST, instance=indirect_costs_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')

        if form.is_valid(): 
            print('form is valid')

            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            indirect_costs = form.save(commit=False)
            indirect_costs.participant_num = participant
            indirect_costs.save()

            # instances = formset.save(commit=False)
            # for instance in instances:
            #     instance.indirect_costs = indirect_costs
            #     instance.save()

            #word_doc = generate_word_document(form.cleaned_data)
            #response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            #response['Content-Disposition'] = f'attachment; filename=indirect_costs.docx'
            #word_doc.save(response)
            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_indirect_costs', args=[participant_id]) + '?submitted=True')
        else:
            print(form.errors)
            #print(formset.errors)
    elif 'save_continue' in request.GET:
        form = Indirect_costs_form(request.POST, instance=indirect_costs_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')
        if form.is_valid():
            indirect_costs = form.save(commit=False)
            indirect_costs.participant_num = participant
            indirect_costs.save()


        return HttpResponseRedirect(reverse('add_stai_form_y', args=[participant_id]))
    else:
        print("Getting Data")
        print("Indirect Costs Data:", indirect_costs_data)

        form = Indirect_costs_form(instance=indirect_costs_data)
        
        if indirect_costs_data:
            print('yes')
            #formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.filter(indirect_costs=indirect_costs_data), prefix='relatives')
        else:
            print('no')
            #formset = RelativeCancerFormSet(queryset=biological_relatives_with_cancer.objects.none(), prefix='relatives')
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form': form,
        'participant_id': participant_id,
        'indirect_costs_data' : indirect_costs_data,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print(form.errors)
    #print(formset.errors)
    #print(request.POST)
    return render(request, 'DataEntry/indirect_costs.html', context)


@login_required(login_url='login')
def add_stai_y1_2(request, participant_id=None):
    print("STAI Form Y-1 & 2 View accessed!")
    if participant_id: 
        participant = get_object_or_404(Participant, participant_number=participant_id)
        stai_y1_2_data = Stai_y1_2.objects.filter(participant_num=participant).first()
        print('the participant number is: ', participant_id)
    
    #RelativeCancerFormSet = modelformset_factory(biological_relatives_with_cancer, form=biological_relatives_with_cancer_form, extra=1)
    
    if request.method == 'POST':
        form = Stai_y1_2_form(request.POST, instance=stai_y1_2_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')

        if form.is_valid(): 
            print('form is valid')

            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            stai_y1_2 = form.save(commit=False)
            stai_y1_2.participant_num = participant
            stai_y1_2.save()

            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_stai_y1_2', args=[participant_id]) + '?submitted=True')
        else:
            print(form.errors)
            #print(formset.errors)
    elif 'save_continue' in request.GET:
        form = Stai_y1_2_form(request.POST, instance=stai_y1_2_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')
        if form.is_valid():
            stai_y1_2 = form.save(commit=False)
            stai_y1_2.participant_num = participant
            stai_y1_2.save()


        return HttpResponseRedirect(reverse('add_stai_y1_2', args=[participant_id]))
    else:
        print("Getting Data")
        print("STAI Form Y-1 & 2 Data:", stai_y1_2_data)

        form = Stai_y1_2_form(instance=stai_y1_2_data)
        
        if stai_y1_2_data:
            print('yes')
        else:
            print('no')
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form': form,
        'participant_id': participant_id,
        'stai_y1_2_data' : stai_y1_2_data,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print(form.errors)
    #print(formset.errors)
    #print(request.POST)
    return render(request, 'DataEntry/stai_y1_2.html', context)


@login_required(login_url='login')
def add_veterans_rand_12(request, participant_id=None):
    print("Veterans Rand 12 View accessed!")
    if participant_id: 
        participant = get_object_or_404(Participant, participant_number=participant_id)
        veterans_rand_12_data = Veterans_rand_12.objects.filter(participant_num=participant).first()
        print('the participant number is: ', participant_id)
    
    #RelativeCancerFormSet = modelformset_factory(biological_relatives_with_cancer, form=biological_relatives_with_cancer_form, extra=1)
    
    if request.method == 'POST':
        form = Veterans_rand_12_form(request.POST, instance=veterans_rand_12_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')

        if form.is_valid(): 
            print('form is valid')

            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            veterans_rand_12 = form.save(commit=False)
            veterans_rand_12.participant_num = participant
            veterans_rand_12.save()

            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_veterans_rand_12', args=[participant_id]) + '?submitted=True')
        else:
            print(form.errors)
            #print(formset.errors)
    elif 'save_continue' in request.GET:
        form = Veterans_rand_12_form(request.POST, instance=veterans_rand_12_data)
        #formset = RelativeCancerFormSet(request.POST, prefix='relatives')
        if form.is_valid():
            veterans_rand_12 = form.save(commit=False)
            veterans_rand_12.participant_num = participant
            veterans_rand_12.save()


        return HttpResponseRedirect(reverse('add_veterans_rand_12', args=[participant_id]))
    else:
        print("Getting Data")
        print("Veterans Rand 12 Data:", veterans_rand_12_data)

        form = Veterans_rand_12_form(instance=veterans_rand_12_data)
        
        if veterans_rand_12_data:
            print('yes')
        else:
            print('no')
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form': form,
        'participant_id': participant_id,
        'veterans_rand_12_data' : veterans_rand_12_data,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print(form.errors)
    #print(formset.errors)
    #print(request.POST)
    return render(request, 'DataEntry/veterans_rand_12.html', context)



@login_required(login_url='login')
def add_mandatory_questionaire_noParticipant(request):

    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')



def generate_mandatory_questionnaire_csv(form, request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.csv'

    writer = csv.writer(response)
    header = [field.label for field in form]
    writer.writerow(header)

    data = [form[field_name].value() for field_name in form.fields]
    writer.writerow(data)

    return response



def generate_mandatory_questionnaire_pdf(pdf_data, request):
    template_path = 'DataEntry/pdf_template.html' 
    template = get_template(template_path)
    context = {'form': pdf_data} 
    html = template.render(context)
    pdf_buffer = BytesIO()

    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire.pdf'
    pdf_buffer.close()
    return response


@login_required(login_url='login')

def add_mandatory_questionaire_c(request, participant_id=None):
    print(request.GET) 
    
    participant = None
    mandatory_questionaire_c_data = None
    print("View accessed!")
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_c_data = Mandatory_questionaire_c.objects.filter(participant_num=participant).first()
        print(participant)
    
    if request.method == 'POST':
        form = Mandatory_questionaire_form_c(request.POST, instance=mandatory_questionaire_c_data)
        if not participant:
            participant = Participant.objects.create(participant_number=participant_id)
            print('created participant')

        if form.is_valid():
            print('form is valid')
            if not participant:
                participant = Participant.objects.create(participant_number=participant_id)
            
            mandatory_questionaire_c = form.save(commit=False)
            mandatory_questionaire_c.participant_num = participant
            mandatory_questionaire_c.save()
            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_mandatory_questionaire_c', args=[participant_id]) + '?submitted=True')
        else:
            print(form.errors)
            
    elif 'generate_pdf' in request.GET:
        print('button pressed')
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data)
        mandatory_questionaire_c_obj = Mandatory_questionaire_c.objects.get(participant_num=participant_id)
        participant_num = mandatory_questionaire_c_obj.participant_num 
        pdf_data = form.instance
        pdf_buffer = generate_mandatory_questionnaire_c_pdf(pdf_data, request)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename={participant_num}_mandatory_questionnaire_c.pdf'
        return response
    
    elif 'generate_csv' in request.GET:
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data) 
        csv_buffer = generate_mandatory_questionnaire_c_csv(form, request)
        mandatory_questionaire_c_obj = Mandatory_questionaire_c.objects.get(participant_num=participant_id)
        participant_num = mandatory_questionaire_c_obj.participant_num 
        response = HttpResponse(csv_buffer.getvalue(), content_type='application/csv')
        response['Content-Disposition'] = f'attachment; filename={participant_num}_mandatory_questionnaire_c.csv'
        return response

    else:
        print("Form is NOT valid!")
        form = Mandatory_questionaire_form_c(instance=mandatory_questionaire_c_data)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }

    return render(request, 'DataEntry/mandatory_questionaire_c.html', context)

def add_mandatory_questionaire_c_noParticipant(request):

    print(request.GET)

    return render(request, 'DataEntry/participant_search.html')

def add_mandatory_questionaire_c_with_id(request, participant_id):
    return add_mandatory_questionaire_c(request, participant_id)


def generate_mandatory_questionnaire_c_pdf(pdf_data, request):
    template_path = 'DataEntry/mandatory_questionaire_c_pdf.html'
    template = get_template(template_path)
    context = {'form': pdf_data} 
    html = template.render(context)
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire.pdf'
    pdf_buffer.close()
    return response

def generate_mandatory_questionnaire_c_csv(form, request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.csv'

    writer = csv.writer(response)

    header = [field.label for field in form]
    writer.writerow(header)

    data = [form[field_name].value() for field_name in form.fields]
    # writer.writerow(data)
    transformed_data = []
    for value in data:
        if value is True:
            transformed_data.append('Yes')
        elif value is False:
            transformed_data.append('No')
        elif value == '':
            transformed_data.append('N/A')
        else:
            transformed_data.append(value)
    writer.writerow(transformed_data)
    return response



@login_required(login_url='login')
def add_mandatory_questionaire_de(request, participant_id=None):
    print(request.POST)
    print(f"URL Parameter - participant_id: {participant_id}")
    print(f"Participant ID: {participant_id}")
    print(f"Request method: {request.method}")
    print("View accessed!")

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        mandatory_questionaire_data = last_mandatory_questionnaire.objects.filter(participant_num=participant).first()
        
        print(participant)

    if request.method == 'POST':
        form = last_mandatory_questionnaire_form(request.POST, instance=mandatory_questionaire_data, initial={'participant_num': participant})
        try:
            form.save()
            if participant_id:
                return HttpResponseRedirect(reverse('add_mandatory_questionaire_de', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_mandatory_questionaire_de') + '?submitted=True')
            print('saved')
        except Exception as e:
            print(f"Error saving forms: {e}")
            print("Form errors:", form.errors)

        messages.success(request, 'Data saved successfully.')
        return HttpResponseRedirect(reverse('add_mandatory_questionaire_de', args=[participant_id]) + '?submitted=True')
    
    elif 'generate_pdf' in request.GET:
        form = last_mandatory_questionnaire_form(instance=mandatory_questionaire_data)

        pdf_buffer = generate_mandatory_questionnaire_d_pdf(form, request)
        mandatory_questionaire_de_obj = last_mandatory_questionnaire.objects.get(participant_num=participant_id)
        participant_num = mandatory_questionaire_de_obj.participant_num 
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename={participant_num}_mandatory_questionnaire_de.pdf'
        return response
    
    elif 'generate_csv' in request.GET:
        form = last_mandatory_questionnaire_form(instance=mandatory_questionaire_data)
        mandatory_questionaire_de_obj = last_mandatory_questionnaire.objects.get(participant_num=participant_id)
        participant_num = mandatory_questionaire_de_obj.participant_num 
        csv_buffer = generate_mandatory_questionnaire_d_csv(form, request)
        response = HttpResponse(csv_buffer.getvalue(), content_type='application/csv')
        response['Content-Disposition'] = f'attachment; filename={participant_num}_mandatory_questionnaire_de.csv'
        return response
    else:
        form = last_mandatory_questionnaire_form(instance=mandatory_questionaire_data)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print(form.errors)

    return render(request, 'DataEntry/mandatory_questionaire_de_new.html', context)


@login_required(login_url='login')
def add_mandatory_questionaire_de_noParticipant(request):

    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def generate_mandatory_questionnaire_d_pdf(form, fg_form, request):
    template_path = 'DataEntry/mandatory_questionaire_d_pdf.html'  
    template = get_template(template_path)
    context = {
        'form': form,
        'fg_form': fg_form,
    } 
    html = template.render(context)
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.pdf'
    pdf_buffer.close()
    return response

def generate_mandatory_questionnaire_d_csv(form, request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.csv'

    writer = csv.writer(response)

    header = [field.label for field in form]
    writer.writerow(header)

    data = [form[field_name].value() for field_name in form.fields]
    # writer.writerow(data)
    transformed_data = []
    for value in data:
        if value is True:
            transformed_data.append('Yes')
        elif value is False:
            transformed_data.append('No')
        elif value == '':
            transformed_data.append('N/A')
        else:
            transformed_data.append(value)

    writer.writerow(transformed_data)
    return response

@login_required(login_url='login')
def add_exposure_form(request, participant_id=None):
    participant = None
    exposure2_data = None

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        exposure2_data = Exposure2.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        print(request.POST)
        form2 = Exposure_Form2(request.POST, instance=exposure2_data, initial={'participant_num': participant})

        if form2.is_valid():
            exposure2_instance = form2.save(commit=False)
            exposure2_instance.participant_num = participant
            exposure2_instance.save()
            
            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_exposure_form', args=[participant_id]) + '?submitted=True')
        else:
            print("Form2 errors:", form2.errors)
    elif 'generate_pdf' in request.GET:
        print('button pressed')
        form2 = Exposure_Form2(instance=exposure2_data)
        pdf_buffer = generate_exposure_questionnaire_pdf(form2, request)
        print(pdf_buffer)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        print(response)
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.pdf'
        print('should be working here')
        return response

    elif 'generate_csv' in request.GET:
        form2 = Exposure_Form2(instance=exposure2_data)
        csv_buffer = generate_exposure_csv(form2, request)
        response = HttpResponse(csv_buffer.getvalue(), content_type='application/csv')
        response['Content-Disposition'] = f'attachment; filename=mandatory_questionnaire.csv'
        return response
    
    else:
        form2 = Exposure_Form2(instance=exposure2_data, initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form2': form2,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
   
    return render(request, 'DataEntry/exposure_form.html', context)

@login_required(login_url='login')
def add_residence_history_form(request, participant_id=None):
    participant = None
    exposure3_data = None

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        exposure3_data = Exposure3.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        form3 = Exposure_Form3(request.POST, instance=exposure3_data, initial={'participant_num': participant})

        if form3.is_valid():
            exposure3_instance = form3.save(commit=False)
            exposure3_instance.participant_num = participant
            exposure3_instance.save()

            messages.success(request, 'Data saved successfully.')
            return HttpResponseRedirect(reverse('add_residence_history_form', args=[participant_id]) + '?submitted=True')
        else:
            print("Form3 errors:", form3.errors)
    elif 'generate_pdf' in request.GET:
        print('button pressed')
        form3 = Exposure_Form3(instance=exposure3_data)
        pdf_buffer = generate_exposure_questionnaire_pdf(form3, request)
        print(pdf_buffer)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        print(response)
        response['Content-Disposition'] = f'attachment; filename=residence_history.pdf'
        print('should be working here')
        return response

    elif 'generate_csv' in request.GET:
        form3 = Exposure_Form3(instance=exposure3_data) 
        csv_buffer = generate_exposure_csv(form3, request)
        response = HttpResponse(csv_buffer.getvalue(), content_type='application/csv')
        response['Content-Disposition'] = f'attachment; filename=residence_history.csv'
        return response
    
    else:
        form3 = Exposure_Form3(instance=exposure3_data, initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form3': form3,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
   
    return render(request, 'DataEntry/residence_history_form.html', context)

def add_exposure_form_noParticipant(request):

    print(request.GET)  

    return render(request, 'DataEntry/participant_search.html')


def add_residence_form_noParticipant(request):

    print(request.GET)  

    return render(request, 'DataEntry/participant_search.html')

def generate_exposure_questionnaire_pdf(form2, form3, request):
    template_path = 'DataEntry/exposure_questionaire_pdf.html'  
    template = get_template(template_path)
    context = {
        'form2': form2,
        'form3': form3,
    }  
    html = template.render(context)
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        return HttpResponse('PDF generation error', content_type='text/plain')
    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.pdf'
    pdf_buffer.close()
    return response


def generate_exposure_csv(form2, form3, request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename=mandatory_questionnaire_d.csv'

    writer = csv.writer(response)

    header = [field.label for field in form2]
    header2 = [field.label for field in form3]
    total_header = header + header2
    writer.writerow(total_header)

    data = [form2[field_name].value() for field_name in form2.fields]
    data2 = [form3[field_name].value() for field_name in form3.fields]
    transformed_data = []
    for value in data + data2:
        if value is True:
            transformed_data.append('Yes')
        elif value is False:
            transformed_data.append('No')
        elif value == '':
            transformed_data.append('N/A')
        else:
            transformed_data.append(value)
    writer.writerow(transformed_data)
    return response



@login_required(login_url='login')
def add_yearly_update(request, participant_id=None):
    participant = None
    annual_study_update_data = None

    if participant_id:
        print('here')
        participant = get_object_or_404(Participant, participant_number=participant_id)
        annual_study_update_data = annual_study_update.objects.filter(participant_num=participant).first()


    if request.method == 'POST':
        form = annual_study_update_form(request.POST, instance=annual_study_update_data, initial={'participant_num': participant})
        
        if form.is_valid():
            annual_study_update_form_instance = form.save(commit=False)
            annual_study_update_form_instance.participant_num = participant
            annual_study_update_form_instance.save()


            if participant_id:
                return HttpResponseRedirect(reverse('add_yearly_update', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_yearly_update') + '?submitted=True')
        else:
            print("Form errors:", form.errors)
    
    else:
        print('here now')
        form = annual_study_update_form(instance=annual_study_update_data, initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
   
    return render(request, 'DataEntry/add_yearly_update.html', context)

@login_required(login_url='login')
def add_yearly_update_noParticipant(request):
    print(request.GET)  

    return render(request, 'DataEntry/participant_search.html')

@login_required(login_url='login')
def add_breath_collection(request, participant_id=None):
    participant = None
    breath_collection_data = None

    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)
        breath_collection_data = Breath_Collection.objects.filter(participant_num=participant).first()

    form = Breath_Collection_Form(instance=breath_collection_data, initial={'participant_num': participant})

    if request.method == 'POST':
        form = Breath_Collection_Form(request.POST, instance=breath_collection_data, initial={'participant_num': participant})

        if form.is_valid():
            form.save()

            if participant_id:
                return HttpResponseRedirect(reverse('add_breath_collection', args=[participant_id]) + '?submitted=True')
            else:
                return HttpResponseRedirect(reverse('add_breath_collection') + '?submitted=True')
            
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print(form.errors)
    print(participant_id)

    return render(request, 'DataEntry/breath_collection.html', context)

def add_breath_collection_noParticipant(request):
    print(request.GET)  # Print the GET parameters

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_ct_scan_dashboard(request, participant_id=None):
    print(request.GET) 
    ct_scans = []
    ct_scans.extend(ct_scan_baseline.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t1.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t2.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t3.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t4.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t5.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t6.objects.filter(participant_num=participant_id))
    ct_scans.extend(ct_scan_t7.objects.filter(participant_num=participant_id))

    ct_scan_baseline_data = ct_scan_baseline.objects.filter(participant_num=participant_id).first()
    ct_scan_t1_data = ct_scan_t1.objects.filter(participant_num=participant_id).first()
    ct_scan_t2_data = ct_scan_t2.objects.filter(participant_num=participant_id).first()
    ct_scan_t3_data = ct_scan_t3.objects.filter(participant_num=participant_id).first()
    ct_scan_t4_data = ct_scan_t4.objects.filter(participant_num=participant_id).first()
    ct_scan_t5_data = ct_scan_t5.objects.filter(participant_num=participant_id).first()
    ct_scan_t6_data = ct_scan_t6.objects.filter(participant_num=participant_id).first()
    ct_scan_t7_data = ct_scan_t7.objects.filter(participant_num=participant_id).first()


    data_present = len(ct_scans) > 0

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scans': ct_scans,
        'participant_id': participant_id,
        'data_present': data_present,
        'ct_scan_baseline_data': ct_scan_baseline_data,
        'ct_scan_t1_data': ct_scan_t1_data,
        'ct_scan_t2_data': ct_scan_t2_data,
        'ct_scan_t3_data': ct_scan_t3_data,
        'ct_scan_t4_data': ct_scan_t4_data,
        'ct_scan_t5_data': ct_scan_t5_data,
        'ct_scan_t6_data': ct_scan_t6_data,
        'ct_scan_t7_data': ct_scan_t7_data,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    
    return render(request, 'DataEntry/ct_scan_dashboard.html', context)

@login_required(login_url='login')
def add_ct_scan_dashboard_no_participant(request):
    print(request.GET) 
    
    return render(request, 'DataEntry/participant_search.html')




@login_required(login_url='login')
def add_ct_scan_baseline(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_baseline.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_Baseline(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_Baseline(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
    if 'generate_pdf' in request.GET:
        ct_scan_form = CT_Scan_Form_Baseline(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        context = {
            'ct_scan_form': ct_scan_form,
            'participant_id': participant_id,
            'ct_nodule_form': ct_nodule_form,
            'ct_nodule_form2': ct_nodule_form2,
            'ct_nodule_form3': ct_nodule_form3,
            'ct_nodule_form4': ct_nodule_form4,
            'ct_nodule_form5': ct_nodule_form5,
        }
        print('button pressed')
        pdf_buffer = download_ct_scan_pdf(request, context)
        response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=CT Scan.pdf'
        return response

    else:
        ct_scan_form = CT_Scan_Form_Baseline(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)



def download_ct_scan_pdf(request, context):
    template = get_template('DataEntry/ct_scan.html')
    html = template.render(context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename="CT Scan.pdf"'

    pisa_status = pisa.CreatePDF(
        html, dest=response, encoding='utf-8')

    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')
    return response



@login_required(login_url='login')
def add_ct_scan_baseline_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')

@login_required(login_url='login')
def add_ct_scan_t1(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t1.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T1(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T1(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T1(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t1_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_ct_scan_t2(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t2.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T2(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T2(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T2(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t2_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_ct_scan_t3(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t3.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T3(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T3(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T3(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t3_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_ct_scan_t4(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t4.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T4(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T4(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T4(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t4_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')

@login_required(login_url='login')
def add_ct_scan_t5(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t5.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T5(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T5(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T5(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t5_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_ct_scan_t6(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t6.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T6(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T6(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T6(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t6_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')

@login_required(login_url='login')
def add_ct_scan_t7(request, participant_id=None):
    print("participant_id received:", participant_id)
    if participant_id:
        participant = get_object_or_404(Participant, participant_number=participant_id)


        ct_scan_data = ct_scan_t7.objects.filter(participant_num=participant).first()
        ct_nodule_data = ct_scan_nodule_1.objects.filter(participant_num=participant).first()
        ct_nodule_data2 = ct_scan_nodule_2.objects.filter(participant_num=participant).first()
        ct_nodule_data3 = ct_scan_nodule_3.objects.filter(participant_num=participant).first()
        ct_nodule_data4 = ct_scan_nodule_4.objects.filter(participant_num=participant).first()
        ct_nodule_data5 = ct_scan_nodule_5.objects.filter(participant_num=participant).first()

    if request.method == 'POST':
        ct_scan_form = CT_Scan_Form_T7(request.POST, instance=ct_scan_data)
        ct_nodule_form = CT_Scan_Nodule_Form_1(request.POST, instance=ct_nodule_data)
        ct_nodule_form2 = CT_Scan_Nodule_Form_2(request.POST, instance=ct_nodule_data2)
        ct_nodule_form3 = CT_Scan_Nodule_Form_3(request.POST, instance=ct_nodule_data3)
        ct_nodule_form4 = CT_Scan_Nodule_Form_4(request.POST, instance=ct_nodule_data4)
        ct_nodule_form5 = CT_Scan_Nodule_Form_5(request.POST, instance=ct_nodule_data5)

        if 'save_ct_scan' in request.POST:
            if ct_scan_form.is_valid():
                ct_scan_data = ct_scan_form.save(commit=False)
                ct_scan_data.participant_num = participant
                ct_scan_data.save()
                messages.success(request, "CT Scan Form saved successfully.")
                ct_scan_form = CT_Scan_Form_T7(instance=ct_scan_data)

        elif 'save_nodule' in request.POST:
            if ct_nodule_form.is_valid():
                ct_nodule_form.save()
                ct_nodule_form2.save()
                ct_nodule_form3.save()
                ct_nodule_form4.save()
                ct_nodule_form5.save()
                messages.success(request, "CT Nodule Form saved successfully.")
                ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
                ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
                ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
                ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
                ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)

    else:
        ct_scan_form = CT_Scan_Form_T7(instance=ct_scan_data)
        if ct_nodule_data:
            ct_nodule_form = CT_Scan_Nodule_Form_1(instance=ct_nodule_data)
        else:
            ct_nodule_form = CT_Scan_Nodule_Form_1(initial={'participant_num': participant})

        if ct_nodule_data2:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(instance=ct_nodule_data2)
        else:
            ct_nodule_form2 = CT_Scan_Nodule_Form_2(initial={'participant_num': participant})

        if ct_nodule_data3:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(instance=ct_nodule_data3)
        else:
            ct_nodule_form3 = CT_Scan_Nodule_Form_3(initial={'participant_num': participant})

        if ct_nodule_data4:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(instance=ct_nodule_data4)
        else:
            ct_nodule_form4 = CT_Scan_Nodule_Form_4(initial={'participant_num': participant})

        if ct_nodule_data5:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(instance=ct_nodule_data5)
        else:
            ct_nodule_form5 = CT_Scan_Nodule_Form_5(initial={'participant_num': participant})

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'ct_scan_form': ct_scan_form,
        'participant_id': participant_id,
        'ct_nodule_form': ct_nodule_form,
        'ct_nodule_form2': ct_nodule_form2,
        'ct_nodule_form3': ct_nodule_form3,
        'ct_nodule_form4': ct_nodule_form4,
        'ct_nodule_form5': ct_nodule_form5,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print("participant:", participant)
    print("CT Scan Form Errors:", ct_scan_form.errors)
    print("CT Nodule Form Errors:", ct_nodule_form.errors)
    return render(request, 'DataEntry/ct_scan.html', context)

@login_required(login_url='login')
def add_ct_scan_t7_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_blood_collection(request, participant_id=None):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')



@login_required(login_url='login')
def add_blood_collection_participant(request, participant_id=None):
    form_data = Blood_Collection.objects.filter(participant_num=participant_id).first()
    if request.method == 'POST':
        form = Blood_Collection_Form(request.POST)
        
        if form.is_valid():
            form.participant_num = participant_id
            form.save()
            print('test')
            return HttpResponseRedirect(reverse('add_blood_collection', args=[participant_id]) + '?submitted=True')
        if 'generate_pdf' in request.POST:
            print('getting here')
            generate_blood_collection_pdf(request, participant_id)
            

        else:
            print('not valid')
    
    else:
        form = Blood_Collection_Form(instance=form_data)
        

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()

    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'      
    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    print(form.errors)

    return render(request, 'DataEntry/blood_collection.html', context)




@login_required(login_url='login')
def add_plco_score(request, participant_id=None):
    participant = get_object_or_404(Participant, participant_number=participant_id)
    print(participant_id)
    if request.method == 'POST':
        form = plco_score_form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('success_page') 
    else:
        form = plco_score_form()

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    return render(request, 'DataEntry/plco_score.html', context)

@login_required(login_url='login')
def add_plco_score_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')


@login_required(login_url='login')
def add_participant(request):
    submitted = False
    if request.method == 'POST':
        form_class = ParticipantForm
        form = ParticipantForm(request.POST)
        if form.is_valid():
            form.save()
            return HttpResponse('/add_participant/?submitted=True?id=' + str(form.instance.id))
    form = ParticipantForm
    context = {
        'form': form,
        'submitted': submitted,
        
               }
    return render(request, 'DataEntry/index.html', {'form': form})

@login_required(login_url='login')
def add_protocol_deviations_participant(request, participant_id=None):
    print('view accessed')
    participant = get_object_or_404(Participant, pk=participant_id)

    ProtocolDeviationsFormSet = modelformset_factory(
        Protocol_Deviations,
        form=ProtocolDeviationsForm,  
        extra=1,
    )
    print('formset factory created')
    if request.method == 'POST':
        print('getting to post')
        print(request.POST)
        formset = ProtocolDeviationsFormSet(request.POST, queryset=Protocol_Deviations.objects.none())
        if formset.is_valid():
            instances = formset.save(commit=False)
            for instance in instances:
                instance.participant_num = participant
                instance.save()
            formset.save_m2m()
            # for form in formset:
            #     form.instance.participant_num = participant
            # print('the formset is valid now')
            # formset.save()
            return render(request, 'DataEntry/protocol_deviations.html', {'formset': formset})
        else:
            print(formset.errors)
    else:
        print('formset invalid')

        # formset = ProtocolDeviationsFormSet(queryset=Protocol_Deviations.objects.none())
        participant_deviations = Protocol_Deviations.objects.filter(participant_num=participant_id)
        formset = ProtocolDeviationsFormSet(queryset=participant_deviations)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'formset': formset,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    return render(request, 'DataEntry/protocol_deviations.html', context)
    # Participant = get_object_or_404(Participant, participant_number=participant_id)
    # form = ProtocolDeviationsForm(request.POST or None, instance=Participant)
    # protocolDeviationFormset = modelformset_factory(Protocol_Deviations, form=ProtocolDeviationsForm, extra=2)
    # formset = protocolDeviationFormset(request.POST or None, queryset=Protocol_Deviations.objects.none())
    # context = {
    #     'form': form,
    #     'formset': formset,
    #     'pariticipant_id': participant_id,
    # }
    # if all([form.is_valid(), formset.is_valid()]):
    #     parent = form.save(commit=False)
    #     parent.save()

    #     for form in formset:
    #         child = form.save(commit=False)
    #         child.parent = parent
    #         child.save()
    #     context['message'] = 'Data saved'
    
    # return render(request, 'add_protocol_deviation', context)


def add_protocol_deviationold2(request):
    ProtocolDeviationsFormSet = formset_factory(ProtocolDeviationsForm, extra=5)
    formset = ProtocolDeviationsFormSet()

    if request.method == 'POST':
        print('post in request')
        formset = ProtocolDeviationsFormSet(request.POST)
        if formset.is_valid():
            print('formset is valid')
            for form in formset:
                if form.cleaned_data:
                    form.save()
            return HttpResponse('/add_protocol_deviation/?submitted=True')

    else:
       
        formset = ProtocolDeviationsFormSet(queryset=Protocol_Deviations.objects.all())
        # existing_data = Protocol_Deviations.objects.all()
        # initial_data = [{'participant_num': item.participant_num.id,  # Use the correct field name for participant_num
        #                  'deviation_type': item.deviation_type,
        #                  'deviation_date': item.deviation_date,
        #                  'clinical_staff_notified': item.clinical_staff_notified,
        #                  'deviation_comments': item.deviation_comments} for item in existing_data]
        # formset = ProtocolDeviationsFormSet(initial=initial_data, prefix='protocol_deviations')

    return render(request, 'add_protocol_deviations.html', {'formset': formset})

def add_protocol_deviation_old(request):
    print('here')
    ProtocolDeviationsFormSet = formset_factory(ProtocolDeviationsForm, extra=0) 
    print('here now')
    if request.method == 'POST':
        for query in connection.queries:
            print(query)
        print('getting here')
        formset = ProtocolDeviationsFormSet(request.POST, prefix='protocol_deviations')
        if formset.is_valid():
            for form in formset:
                if form.cleaned_data:
                    form.save()
                    print('data saved')
            return redirect('add_protocol_deviations.html')  

    else:
        print('trying to save')
        print(formset.errors)
        formset = ProtocolDeviationsFormSet(prefix='protocol_deviations')

    return render(request, 'add_protocol_deviations.html', {'formset': formset})

    # submitted = False
    # if request.method == 'POST':
    #     form_class = ParticipantForm
    #     form = ParticipantForm(request.POST)
    #     if form.is_valid():
    #         form.save()
    #         return HttpResponse('/add_participant/?submitted=True')
    # form = ParticipantForm
    # return render(request, 'DataEntry/index.html', {'form': form})




@login_required(login_url='login')
def add_protocol_deviations_noParticipant(request):
    print(request.GET) 

    return render(request, 'DataEntry/participant_search.html')




@login_required(login_url='login')
def update_participant_status(request):
    if request.method == 'POST':
        inclusion_criteria_values = [
            request.POST.get('inclusion_criteria_1'),
            request.POST.get('inclusion_criteria_2'),
        ]

        if all(value == 'yes' for value in inclusion_criteria_values[:5]) and all(value == 'no' for value in inclusion_criteria_values[5:]):
            participant_status = 'group 1'
        else:
            participant_status = 'group 2'

        inclusion_object = inclusion.objects.first() 
        inclusion_object.participant_status = participant_status
        inclusion_object.save()

        response_data = {'participant_status': participant_status}
        return JsonResponse(response_data)
    else:
        return render(request, 'form.html')


@login_required(login_url='login')
def add_lab_processing(request):
    if request.method == 'POST':
        form = Lab_Processing_Form(request.POST, request.FILES)
        if form.is_valid():
            lab_processing_obj = lab_processing()
            lab_processing_obj.lab_processing_upload = form.cleaned_data['lab_processing_upload']
            lab_processing_obj.save()
    else:
        form = Lab_Processing_Form()

    uploaded_files = lab_processing.objects.all()  

    context = {
        'form': form,
        'uploaded_files': uploaded_files,

    }

    return render(request, 'DataEntry/lab_processing.html', context)



# @login_required(login_url='login')
# def add_lab_processing_participant(request, participant_id=None):
#     if request.method == 'POST':
#         if 'upload' in request.POST:  
#             form = UploadFileForm(request.POST, request.FILES)
#             if form.is_valid():
#                 form.save()
#                 return redirect('manage_files')
#         elif 'delete' in request.POST:  
#             file_id = request.POST.get('delete')
#             file = Participant.objects.get(id=file_id)
#             file.delete()
#             return redirect('manage_files')
#     else:
#         form = UploadFileForm()
#     files = ParticipantFile.objects.all()
#     return render(request, 'myapp/manage_files.html', {'form': form, 'files': files})

    uploaded_files = lab_processing.objects.all()  

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'uploaded_files': uploaded_files,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }

    return render(request, 'DataEntry/lab_processing.html', context)


@login_required(login_url='login')
def add_lab_processing_participant_older(request, participant_id=None):
    if participant_id:
        print('here')
    if request.method == 'POST':
        form = Lab_Processing_Form(request.POST, request.FILES)
        if form.is_valid():
            lab_processing_obj = lab_processing()
            lab_processing_obj.lab_processing_upload = form.cleaned_data['lab_processing_upload']
            lab_processing_obj.save()
    else:
        form = Lab_Processing_Form()

    uploaded_files = lab_processing.objects.all()  

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'uploaded_files': uploaded_files,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }

    return render(request, 'DataEntry/lab_processing.html', context)

from .models import lab_processing
import traceback

@login_required(login_url='login')
def add_lab_processing_participant(request, participant_id=None):
    participant = get_object_or_404(Participant, pk=participant_id)
    
    if request.method == 'POST':
        form = Lab_Processing_Form(request.POST, request.FILES)
        if form.is_valid():
            try:
                lab_processing_obj = form.save(commit=False)
                lab_processing_obj.participant_num = participant
                lab_processing_obj.save()
                return redirect('add_lab_processing_participant', participant_id=participant_id)
            except Exception as e:
                print("Error saving file:", e)
                traceback.print_exc()
        else:
            print("Form errors:", form.errors)
    else:
        form = Lab_Processing_Form()

    uploaded_files = lab_processing.objects.filter(participant_num=participant)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant).first()
    date_of_birth = base_data.date_of_birth if base_data else 'N/A'
    initials = base_data.initials if base_data else 'N/A'

    context = {
        'form': form,
        'uploaded_files': uploaded_files,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }

    return render(request, 'DataEntry/lab_processing.html', context)

def delete_lab_processing_file(request, file_id):
    if request.method == 'POST':
        file_to_delete = get_object_or_404(lab_processing, pk=file_id)
        participant_id = file_to_delete.participant_num_id  # Retrieve participant ID before deleting
        file_to_delete.delete()
    return redirect('add_lab_processing_participant', participant_id=participant_id)



def delete_file(request, file_id):
    file_to_delete = get_object_or_404(lab_processing, id=file_id)
    file_to_delete.delete()

    return redirect('DataEntry/lab_processing.html')

@login_required(login_url='login')
def update_mandatory_questionaire(request, participant_id):
    mandatory_questionaire = Mandatory_questionaire.objects.get(pk=participant_id)
    form = Mandatory_questionaire_form(request.POST or None, instance=mandatory_questionaire)
    return render(request, 'DataEntry/update_mandatory_questionaire.html', {'Mandatory_questionaire': Mandatory_questionaire, 'participant_id': participant_id})

@login_required(login_url='login')
def add_clinical_procedures(request):
    return render(request, 'DataEntry/participant_search.html', {})



@login_required(login_url='login')
def add_clinical_procedures_participant(request, participant_id=None):
    participant = get_object_or_404(Participant, pk=participant_id)
    ClinicalProcedureFormSet = modelformset_factory(
        Clinical_Procedures,
        form=clinical_procedures_form,  
        extra=1,
    )

    if request.method == 'POST':
        formset = ClinicalProcedureFormSet(request.POST, request.FILES)
        if formset.is_valid():
            instances = formset.save(commit=False)
            for instance in instances:
                instance.participant_num = participant
                instance.save()
            formset.save_m2m()
            return render(request, 'DataEntry/clinical_procedures.html', {'formset': formset})
        else:
            print(formset.errors)
    else:
         
        clinical_procedures = Clinical_Procedures.objects.filter(participant_num=participant)
        formset = ClinicalProcedureFormSet(queryset=clinical_procedures)

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'formset': formset,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    return render(request, 'DataEntry/clinical_procedures.html', context)



@login_required(login_url='login')
def add_clinical_procedures_participant_oldFeb8(request, participant_id=None):
    print('view accessed')
    
    participant = get_object_or_404(Participant, pk=participant_id)
    if participant_id:
        participant_num = participant_id
        
    ClinicalProcedureFormSet = modelformset_factory(
        Clinical_Procedures,
        form=clinical_procedures_form,  
        extra=1,
    )
    print('formset factory created')
    if request.method == 'POST':
        print('getting to post')
        print(request.POST)
        participant = get_object_or_404(Participant, pk=participant_id)
        #formset = ClinicalProcedureFormSet(request.POST, queryset=Clinical_Procedures.objects.none())
        formset = ClinicalProcedureFormSet(request.POST, request.FILES, queryset=Clinical_Procedures.objects.filter(participant_num=participant))
        if formset.is_valid():
            for form in formset:
                form.instance.participant_num = participant
            print('the formset is valid now')
            formset.save(commit=False)
            formset.partician_num = participant_id
            formset.save()
            return render(request, 'DataEntry/clinical_procedures.html', {'formset': formset})
        else:
            print(formset.errors)
    else:
        print('formset invalid')

        #formset = ClinicalProcedureFormSet(queryset=Clinical_Procedures.objects.none())
        formset = ClinicalProcedureFormSet(queryset=Clinical_Procedures.objects.filter(participant_num=participant))
        clinical_procedures = Clinical_Procedures.objects.filter(participant_num=participant_id)
        formset = ClinicalProcedureFormSet(queryset=clinical_procedures)
    context = {
        'formset': formset,
        'participant_id': participant_id,
    }
    return render(request, 'DataEntry/clinical_procedures.html', context)

@login_required(login_url='login')
def add_protocol_deviations(request):
       return render(request, 'DataEntry/participant_search.html', {})


# @login_required(login_url='login')
# def add_protocol_deviations_participant(request, participant_id=None):
       
#     context = {
#         'participant_id': participant_id,
#     }

#     return render(request, 'DataEntry/protocol_deviations.html', context)

@login_required(login_url='login')
def add_plco_score_participant(request, participant_id=None):
    if request.method == 'POST':
        form = plco_score_form(request.POST)
        if form.is_valid():
            form.save()
            return redirect('success_page') 
    else:
        form = plco_score_form()

    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'

    context = {
        'form': form,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    return render(request, 'DataEntry/plco_score.html', context)

@login_required
def add_plco_score(request):

    return render(request, 'DataEntry/participant_search.html', {})

@login_required(login_url='login')
def add_lab_processing_with_data(request, participant_num):
    participant = lab_processing.objects.get(participant_num=participant_num)
    form = Lab_Processing_Form(instance=participant)

    return render(request, 'lab_processing_with_data.html', {'form': form})

@login_required(login_url='login')
def add_participant(request):
    if request.method == 'POST':
        form = ParticipantForm(request.POST)
        if form.is_valid():
            form.save()
            return JsonResponse({'message': 'Participant form submitted successfully'})
        else:
            return JsonResponse({'error': form.errors}, status=400)



@login_required(login_url='login')
def search(request):
    query = request.GET.get('q')
    if query:
        participant = get_object_or_404(Participant, participant_number=query)
        return redirect('participant_detail', participant_id=participant.participant_number)
    else:
        return render(request, 'search.html')

@login_required(login_url='login')
def list_participants(request):
    participants = Participant.objects.all()
    context = {
        'participants': participants,
    }
    return render(request,'DataEntry/list_participants.html', context)


@login_required(login_url='login')
def show_participants(request, participant_id):
    participants = Participant.objects.get(pk=participant_id)
    context = {
        'participants': participants,
    }
    return render(request,'DataEntry/show_participants.html', context)

@login_required(login_url='login')
def history_page(request):
    histories = History.objects.all().order_by('-timestamp')
    return render(request, 'history.html', {'histories': histories})

@login_required(login_url='login')
def add_data_test(request):
    return render(request, 'DataEntry/add_data.html', {})






@login_required(login_url='login')
def dashboard(request):
    participants = Participant.objects.all()
    data = []

    for participant in participants:
        participant_data = {
            'participant': participant,
            'inclusion_data': inclusion.objects.filter(participant_num=participant).exists(),
            'questionnaire_data': Mandatory_questionaire.objects.filter(participant_num=participant).exists(),
            'questionnaire_c_data': Mandatory_questionaire_c.objects.filter(participant_num=participant).exists(),
            'questionnaire_d_data': Mandatory_questionaire_de.objects.filter(participant_num=participant).exists(),
            'exposure_data': Exposure.objects.filter(participant_num=participant).exists(),
            'breath_collection_data': Breath_Collection.objects.filter(participant_num=participant).exists(),
            'blood_collection_data': Blood_Collection.objects.filter(participant_num=participant).exists(),
            'ct_scan_data': ct_scan_baseline.objects.filter(participant_num=participant).exists(),
            'yearly_update_data': annual_study_update.objects.filter(participant_num=participant).exists(),

        }
        data.append(participant_data)


    context = {
        'data': data,

        }
    return render(request, 'dashboard.html', context)



@login_required(login_url='login')
def data_download(request):
    participants = Participant.objects.all()
    data = []

    for participant in participants:
        participant_data = {
            'participant': participant,
            'inclusion_data': inclusion.objects.filter(participant_num=participant).exists(),
            'questionnaire_data': Mandatory_questionaire.objects.filter(participant_num=participant).exists(),
            'questionnaire_c_data': Mandatory_questionaire_c.objects.filter(participant_num=participant).exists(),
            'questionnaire_d_data': Mandatory_questionaire_de.objects.filter(participant_num=participant).exists(),
            'exposure_data': Exposure.objects.filter(participant_num=participant).exists(),
            'breath_collection_data': Breath_Collection.objects.filter(participant_num=participant).exists(),
            'blood_collection_data': Blood_Collection.objects.filter(participant_num=participant).exists(),
            'ct_scan_data': ct_scan_baseline.objects.filter(participant_num=participant).exists(),
            'yearly_update_data': annual_study_update.objects.filter(participant_num=participant).exists(),

        }
        data.append(participant_data)

    context = {'data': data}
    return render(request, 'dashboard.html', context)





@login_required(login_url='login')
def dashboard_participant(request, participant_id=None):
    participants = Participant.objects.all()
    data = []
    if participant_id:
        print('participant_id:', participant_id)
    for participant in participants:
        participant_data = {
            'participant': participant,
            'inclusion_data': inclusion.objects.filter(participant_num=participant).exists(),
            'questionnaire_data': Mandatory_questionaire.objects.filter(participant_num=participant).exists(),
            'questionnaire_c_data': Mandatory_questionaire_c.objects.filter(participant_num=participant).exists(),
            'questionnaire_d_data': Mandatory_questionaire_de.objects.filter(participant_num=participant).exists(),
            'exposure_data': Exposure.objects.filter(participant_num=participant).exists(),
            'breath_collection_data': Breath_Collection.objects.filter(participant_num=participant).exists(),
            'blood_collection_data': Blood_Collection.objects.filter(participant_num=participant).exists(),
            'ct_scan_data': ct_scan_baseline.objects.filter(participant_num=participant).exists(),
            'yearly_update_data': annual_study_update.objects.filter(participant_num=participant).exists(),

        }
        data.append(participant_data)
    base_data = Mandatory_questionaire.objects.filter(participant_num=participant_id).first()
    if base_data:
        date_of_birth = base_data.date_of_birth
        initials = base_data.initials

    else:
        date_of_birth = 'N/A'
        initials = 'N/A'
    
    context = {
        'data': data,
        'participant_id': participant_id,
        'initials': initials,
        'date_of_birth': date_of_birth,
    }
    return render(request, 'dashboard.html', context)

@login_required(login_url='login')
def check_participant(request, number):
    try:
        participant = Participant.objects.get(participant_number=number)
        return JsonResponse({'exists': True})
    except Participant.DoesNotExist:
        return JsonResponse({'exists': False, 'message': 'Participant not found in the database.'})
    except Exception as e:
        return JsonResponse({'exists': False, 'message': str(e)})
    



def concatenate_runs(paragraph):
    """
    Concatenate the text from consecutive runs within a paragraph.
    """
    concatenated_text = ''
    for run in paragraph.runs:
        concatenated_text += run.text
    return concatenated_text

def replace_placeholders_in_paragraph(paragraph, replacements):
    """
    Replace placeholders within a paragraph.
    """
    for placeholder, replacement in replacements.items():
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)


def count_tables_and_cells(doc):
    num_tables = len(doc.tables)
    num_cells = sum(len(table._cells) for table in doc.tables)
    return num_tables, num_cells



def search_and_replace_table_values(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for placeholder, replacement in replacements.items():
                        run.text = run.text.replace(placeholder, replacement)

def generate_blood_collection_pdf(request, participant_id):
    print(participant_id)
    blood_collection_data = Blood_Collection.objects.get(participant_num=participant_id)
    mandatory_questionaire_data = Mandatory_questionaire.objects.get(participant_num=participant_id)

    template_path = os.path.join(settings.BASE_DIR, 'IDEAL_Questionnaire_Version_30-May-2023_CB- FINAL V2.docx')
    doc = Document(template_path)
    print(blood_collection_data.participant_num)

    num_tables, num_cells = count_tables_and_cells(doc)
    print(f"Number of tables: {num_tables}")
    print(f"Number of cells: {num_cells}")
    
    replacements = {



        '{postal_code}' : str(mandatory_questionaire_data.postal_code),
        '{date_of_birth}' : str(mandatory_questionaire_data.date_of_birth),
        '{current_age} ' : str(mandatory_questionaire_data.current_age),
        '{current_height_metric}' : str(mandatory_questionaire_data.current_height_metric),
        '{current_height_feet}' : str(mandatory_questionaire_data.current_height_feet),
        '{current_height_inches}' : str(mandatory_questionaire_data.current_height_inches),
        '{current_weight} ' : str(mandatory_questionaire_data.current_weight),
        '{current_weight_unit}' : str(mandatory_questionaire_data.current_weight_unit),
        '{sex_birth}' : str(mandatory_questionaire_data.sex_birth),
        '{gender_identity}' : str(mandatory_questionaire_data.gender_identity),
        '{ethnicity}' : str(mandatory_questionaire_data.ethnicity),
        '{ethnicity_other}' : str(mandatory_questionaire_data.ethnicity_other),
        '{born_in_canada}' : str(mandatory_questionaire_data.born_in_canada),
        '{year_moved_to_canada}' : str(mandatory_questionaire_data.year_moved_to_canada),
        '{birthplace}' : str(mandatory_questionaire_data.birthplace),
        '{highest_education_lvl}' : str(mandatory_questionaire_data.highest_education_lvl),
        '{highest_education_lvl_other}' : str(mandatory_questionaire_data.highest_education_lvl_other),
        '{copd}' : str(mandatory_questionaire_data.copd),
        '{emphysema}': str(mandatory_questionaire_data.emphysema),
        '{chronic_bronchitis}' : str(mandatory_questionaire_data.chronic_bronchitis),
        '{asthma}' : str(mandatory_questionaire_data.asthma),
        '{diabetes}' : str(mandatory_questionaire_data.diabetes),
        '{hypertension}' : str(mandatory_questionaire_data.hypertension),
        '{tuberculosis}' : str(mandatory_questionaire_data.tuberculosis),
        '{adult_pneumonia}' : str(mandatory_questionaire_data.adult_pneumonia),
        '{pulmonary_fibrosis}' : str(mandatory_questionaire_data.pulmonary_fibrosis),
        '{hiv}' : str(mandatory_questionaire_data.hiv),
        '{long_covid}' : str(mandatory_questionaire_data.long_covid),
        '{personal_cancer_history}  ' : str(mandatory_questionaire_data.personal_cancer_history),
        '{personal_cancer_history_youngest_age}': str(mandatory_questionaire_data.personal_cancer_history_youngest_age),
        '{personal_history_cancer_type}' : str(mandatory_questionaire_data.personal_history_cancer_type),


    }

  
    replaced_placeholders = set()  
    for paragraph in doc.paragraphs:
        
        paragraph_text = concatenate_runs(paragraph)
       
        for placeholder, replacement in replacements.items():
            if placeholder not in replaced_placeholders:
                paragraph_text = paragraph_text.replace(placeholder, replacement)
                replaced_placeholders.add(placeholder)
        
        for run in paragraph.runs:
            run.text = paragraph_text

 
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                
                    paragraph_text = concatenate_runs(paragraph)

                    for placeholder, replacement in replacements.items():
                        if placeholder not in replaced_placeholders:
                            paragraph_text = paragraph_text.replace(placeholder, replacement)
                            replaced_placeholders.add(placeholder)

                    for run in paragraph.runs:
                        run.text = paragraph_text


    
    for table_index, table in enumerate(doc.tables):
        print(f"Table {table_index + 1}:")
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                print(f"  Row {row_index + 1}, Cell {cell_index + 1}:")
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    for run_index, run in enumerate(paragraph.runs):
                        print(f"    Paragraph {paragraph_index + 1}, Run {run_index + 1}:")
                        print("      Text:", run.text)

    filled_doc_path = os.path.join(settings.BASE_DIR, 'filled_doc.docx')
    doc.save(filled_doc_path)

    pdf_path = os.path.join(settings.BASE_DIR, 'generated_pdf.pdf')
    docx2pdf_convert(filled_doc_path, pdf_path)


    with open(pdf_path, 'rb') as pdf_file:
        pdf_data = pdf_file.read()

    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="generated_pdf.pdf"'

    return response









from django.shortcuts import render, redirect
from django.views.generic import TemplateView, ListView, CreateView
from django.core.files.storage import FileSystemStorage
from django.urls import reverse_lazy

from .forms import pdf_uploads_questionnaire_form
from .models import pdf_uploads_questionnaire, pdf_uploads






class Home(TemplateView):
    template_name = 'home.html'


def upload(request):
    context = {}
    if request.method == 'POST':
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        name = fs.save(uploaded_file.name, uploaded_file)
        context['url'] = fs.url(name)
    return render(request, 'upload.html', context)


def pdf_list(request):
    books = pdf_uploads_questionnaire.objects.all()
    return render(request, 'book_list.html', {
        'books': books
    })


def upload_pdf(request):
    if request.method == 'POST':
        form = pdf_uploads_questionnaire_form(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('book_list')
    else:
        form = pdf_uploads_questionnaire_form()
    return render(request, 'upload_book.html', {
        'form': form
    })


def delete_pdf(request, participant_id):
    if request.method == 'POST':
        book = pdf_uploads.objects.get(participant_num=participant_id)
        book.delete()
    return redirect('book_list')


class PdfListView(ListView):
    model = pdf_uploads_questionnaire
    template_name = 'class_book_list.html'
    context_object_name = 'books'


class UploadPdfView(CreateView):
    model = pdf_uploads_questionnaire
    form_class = pdf_uploads_questionnaire_form
    success_url = reverse_lazy('class_book_list')
    template_name = 'upload_book.html'




